"""
Shared utility functions for generating Word documents from markdown-like text.
"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_text(paragraph, text, bold=False, italic=False, size=None, color=None, font_name=None):
    """Add a run with specific formatting."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def parse_inline_formatting(paragraph, text, default_size=11, default_font=None):
    """Parse bold (**text**) and code (`text`) in inline text."""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            add_formatted_text(paragraph, inner, bold=True, size=default_size, font_name=default_font)
        elif part.startswith('`') and part.endswith('`'):
            inner = part[1:-1]
            run = paragraph.add_run(inner)
            run.font.size = Pt(default_size - 1)
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        else:
            if part:
                add_formatted_text(paragraph, part, size=default_size, font_name=default_font)


def set_paragraph_spacing(paragraph, before=0, after=6, line_spacing=1.15):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def add_table_from_rows(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parse_inline_formatting(p, h, default_size=10)
        for run in p.runs:
            run.bold = True
        set_cell_shading(cell, '2F5496')
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            parse_inline_formatting(p, val.strip(), default_size=10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'D6E4F0')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph('')  # spacing after table
    return table


def parse_markdown_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (headers, rows, end_idx)."""
    headers = [c.strip() for c in lines[start_idx].strip().strip('|').split('|')]
    data_start = start_idx + 2
    rows = []
    idx = data_start
    while idx < len(lines) and '|' in lines[idx] and lines[idx].strip().startswith('|'):
        row = [c.strip() for c in lines[idx].strip().strip('|').split('|')]
        rows.append(row)
        idx += 1
    return headers, rows, idx


def build_document(doc, lines, default_font='Calibri', cjk_font=None):
    """Build a Word document from parsed lines."""
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == '---':
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith('#'):
            level = len(stripped.split(' ')[0])
            text = stripped.lstrip('#').strip()

            if level == 1:
                p = doc.add_heading(text, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                p = doc.add_heading(text, level=1)
            elif level == 3:
                p = doc.add_heading(text, level=2)
            elif level == 4:
                p = doc.add_heading(text, level=3)
            else:
                p = doc.add_heading(text, level=4)
            i += 1
            continue

        # Tables
        if '|' in stripped and stripped.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            headers, rows, end_idx = parse_markdown_table(lines, i)
            add_table_from_rows(doc, headers, rows)
            i = end_idx
            continue

        # Code blocks
        if stripped.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1

            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=6, after=6)
            for cl in code_lines:
                run = p.add_run(cl + '\n')
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            pPr = p._element.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), 'F5F5F5')
            pPr.append(shd)
            continue

        # Block quotes
        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            parse_inline_formatting(p, text, default_size=11)
            for run in p.runs:
                run.italic = True
            set_paragraph_spacing(p, before=6, after=6)
            i += 1
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.text = ''
            parse_inline_formatting(p, text, default_size=11)
            set_paragraph_spacing(p, before=2, after=2)
            i += 1
            continue

        # Numbered lists
        m = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if m:
            text = m.group(2)
            p = doc.add_paragraph(style='List Number')
            p.text = ''
            parse_inline_formatting(p, text, default_size=11)
            set_paragraph_spacing(p, before=2, after=2)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        parse_inline_formatting(p, stripped, default_size=11)
        set_paragraph_spacing(p, before=3, after=6)
        i += 1

    return doc
