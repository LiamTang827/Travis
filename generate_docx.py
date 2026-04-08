"""
Generate Chinese and English Word documents from the midterm report markdown.
"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_text(paragraph, text, bold=False, italic=False, size=None, color=None, font_name=None):
    """Add a run with specific formatting."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def parse_inline_formatting(paragraph, text, default_size=11, default_font=None):
    """Parse bold (**text**) and code (`text`) in inline text."""
    # Split by ** and ` patterns
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            add_formatted_text(paragraph, inner, bold=True, size=default_size, font_name=default_font)
        elif part.startswith('`') and part.endswith('`'):
            inner = part[1:-1]
            run = paragraph.add_run(inner)
            run.font.size = Pt(default_size - 1)
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        else:
            if part:
                add_formatted_text(paragraph, part, size=default_size, font_name=default_font)


def set_paragraph_spacing(paragraph, before=0, after=6, line_spacing=1.15):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def add_table_from_rows(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parse_inline_formatting(p, h, default_size=10)
        for run in p.runs:
            run.bold = True
        set_cell_shading(cell, '2F5496')
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            parse_inline_formatting(p, val.strip(), default_size=10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'D6E4F0')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph('')  # spacing after table
    return table


def parse_markdown_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (headers, rows, end_idx)."""
    headers = [c.strip() for c in lines[start_idx].strip().strip('|').split('|')]
    # Skip separator line
    data_start = start_idx + 2
    rows = []
    idx = data_start
    while idx < len(lines) and '|' in lines[idx] and lines[idx].strip().startswith('|'):
        row = [c.strip() for c in lines[idx].strip().strip('|').split('|')]
        rows.append(row)
        idx += 1
    return headers, rows, idx


def build_document(doc, lines, default_font='Calibri', cjk_font=None):
    """Build a Word document from parsed lines."""
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip horizontal rules
        if stripped == '---':
            i += 1
            continue

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith('#'):
            level = len(stripped.split(' ')[0])  # count #'s
            text = stripped.lstrip('#').strip()

            if level == 1:
                p = doc.add_heading(text, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                p = doc.add_heading(text, level=1)
            elif level == 3:
                p = doc.add_heading(text, level=2)
            elif level == 4:
                p = doc.add_heading(text, level=3)
            else:
                p = doc.add_heading(text, level=4)
            i += 1
            continue

        # Tables
        if '|' in stripped and stripped.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            headers, rows, end_idx = parse_markdown_table(lines, i)
            add_table_from_rows(doc, headers, rows)
            i = end_idx
            continue

        # Code blocks
        if stripped.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```

            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=6, after=6)
            for cl in code_lines:
                run = p.add_run(cl + '\n')
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # Add shading to the paragraph
            pPr = p._element.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), 'F5F5F5')
            pPr.append(shd)
            continue

        # Block quotes
        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            parse_inline_formatting(p, text, default_size=11)
            for run in p.runs:
                run.italic = True
            set_paragraph_spacing(p, before=6, after=6)
            i += 1
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.text = ''
            parse_inline_formatting(p, text, default_size=11)
            set_paragraph_spacing(p, before=2, after=2)
            i += 1
            continue

        # Numbered lists
        m = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if m:
            text = m.group(2)
            p = doc.add_paragraph(style='List Number')
            p.text = ''
            parse_inline_formatting(p, text, default_size=11)
            set_paragraph_spacing(p, before=2, after=2)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        parse_inline_formatting(p, stripped, default_size=11)
        set_paragraph_spacing(p, before=3, after=6)
        i += 1

    return doc


def create_chinese_doc(md_lines):
    """Create the Chinese version."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # Heading styles
    for i in range(5):
        hs = doc.styles[f'Heading {i+1}'] if i > 0 else doc.styles['Title']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        if hasattr(hs.element, 'rPr') and hs.element.rPr is not None:
            hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    build_document(doc, md_lines)

    path = '/Users/tangliam/CriptoAnalyst/midterm_report_CN.docx'
    doc.save(path)
    print(f'Chinese version saved: {path}')
    return path


# ── English version (restructured per Interim Report requirements) ──

EN_LINES = r"""# Cross-Chain Tracing-Based Cryptocurrency Anti-Money Laundering Risk Identification System
## 1st Interim Report

---

## 1. Introduction

### 1.1 Background

### 1.1 Scale and Evolution of Cryptocurrency Money Laundering

The decentralized, pseudonymous, and globally accessible nature of blockchain has made it a significant channel for illicit fund flows alongside financial innovation. According to Chainalysis's *2024 Crypto Money Laundering Report*, illegal cryptocurrency addresses received approximately **$40.9 billion** in 2023, a figure that grew to over **$154 billion** by 2025 — a 162% year-over-year increase. Notably, stablecoins now account for **63%** of illicit transactions (reaching 84% in 2025), as criminals increasingly shift from Bitcoin to stablecoins like USDT due to their high liquidity, ease of cross-chain transfer, and larger regulatory blind spots.

In terms of crime typology, money laundering techniques have exhibited a systematic trend toward "professionalization." The Huione Group, for example, has processed over **$70 billion** in cryptocurrency transactions since 2021, evolving into an underground financial infrastructure serving the entire fraud and laundering pipeline. This trend indicates that cryptocurrency crime is no longer scattered individual behavior but organized crime with institutional structure and technical barriers.

### 1.2 The Rise of Cross-Chain Bridges and Regulatory Blind Spots

Cross-chain bridges have expanded dramatically in recent years. Stargate Finance alone handles over **$2.3 billion** in monthly cross-chain volume, with the broader DeFi ecosystem facilitating over **$8 billion** in monthly cross-chain asset transfers. While bridges serve a legitimate and critical infrastructure function — enabling asset transfers between different blockchains — this capability has been systematically exploited by criminals to sever fund tracing chains.

Elliptic's 2023 report identified **$7 billion** in illicit assets laundered through cross-chain services, with the figure growing rapidly since 2022. Among identifiable laundering schemes, **58% utilized cross-chain bridges** as a key component (2024 data). Chainalysis similarly reported a significant surge in bridge usage from stolen-fund-linked addresses in 2023.

The operations of North Korean hacking group Lazarus Group illustrate this pattern:
- **Ronin Bridge Attack (March 2022)**: Approximately $625 million stolen, subsequently laundered through Tornado Cash mixing, Avalanche bridge to Bitcoin network, and Sinbad mixer for secondary washing — spanning over 12,000 addresses across multiple chains.
- **Harmony Horizon Bridge Attack (June 2022)**: Approximately $100 million stolen, with **98% of stolen assets processed through Tornado Cash**, followed by repeated hops between Ethereum, BNB Chain, and BitTorrent Chain until partial funds resurfaced on Avalanche and TRON in 2023.

These attacks demonstrate the typical modern laundering pattern: **mixing + cross-chain + multi-hop relay**, designed to exhaust law enforcement and compliance resources through increased tracing difficulty.

### 1.3 Limitations of Existing Tools

Investigations by INTERPOL and related law enforcement agencies show that **74% of agencies report significant limitations in existing blockchain investigation tools for cross-chain activity tracing**. Mainstream tools (Chainalysis Reactor, TRM Labs) possess certain cross-chain capabilities, but their core algorithms have not been publicly validated by academia, and they primarily serve commercial clients.

From the academic research perspective, existing literature has the following main limitations:

- **Single-chain focus**: The vast majority of AML research targets Bitcoin or Ethereum as a single study object, lacking analytical frameworks for cross-chain scenarios.
- **Insufficient blacklist coverage**: Most studies rely on OFAC sanctions lists, overlooking the actual freeze lists maintained by stablecoin issuers (such as Tether) — which are closer to the front line of real-world money laundering detection.
- **No traceability classification**: Existing research rarely distinguishes between "transparent bridges" and "opaque bridges," yet this classification is critical for determining whether fund flow tracing can continue.
- **Lack of depth control in graph analysis**: Current graph analysis methods typically set fixed depths without considering the varying suspicion levels across different branches, leading to either too-shallow tracing (evadable) or too-deep tracing (false positives on normal users).

---

## 2. Research Motivation

### 2.1 Why This Work Matters

The core motivation stems from a practical contradiction: **every transaction on blockchain is publicly auditable, yet fund flows can still be effectively concealed.**

Blockchain's transparency is its key differentiator from traditional banking — all transaction records are permanently stored on a public ledger accessible to anyone. However, criminals neutralize this transparency through three types of tools:

- **Mixers**: Combine funds from multiple users before outputting, severing the input-output correspondence. Tornado Cash anonymized over **$7 billion** in fund flows before being sanctioned by OFAC.
- **Opaque cross-chain bridges**: Transfer assets through liquidity pools (Synapse, Multichain) or market-maker models (Orbiter Finance, Owlto), where on-chain data contains no record of "who the money actually went to."
- **Clean address relays (Money Mules)**: Use a chain of ostensibly "clean" intermediate addresses to increase the graph distance between source and destination addresses, diluting the association.

Therefore, simple one-hop checks of "is this address on a blacklist" are insufficient. Real-world laundering paths often require multi-hop tracing to uncover high-risk associations hidden in intermediate layers.

### 2.2 The Value and Limitations of Stablecoin Blacklists

Tether (USDT issuer) is currently one of the most effective "first responders" in the cryptocurrency space. Its blacklist (USDT Blacklist) has frozen over **8,500 addresses** with assets exceeding **$4.2 billion** as of March 2025, including addresses associated with sanctioned entities, fraud networks, and terrorism financing.

BlockSec's analysis in *Following the Frozen: An On-Chain Analysis of USDT Blacklisting* found that **54% of frozen addresses had already transferred their assets before the freeze occurred**, indicating that freezing actions often lag behind actual fund movements. This means: **the real value lies in tracing fund flows before the freeze event, not merely analyzing the post-freeze snapshot.**

This research builds on this observation: using the USDT blacklist as an "anchor point," tracing upstream addresses that have had fund interactions with blacklisted addresses, and extending through cross-chain tracing to other chains to build a multi-hop, multi-chain risk assessment system.

### 2.3 The Compliance Dilemma for Ordinary Users: Passive Taint

Existing AML tools (Chainalysis, TRM Labs, Elliptic) primarily serve institutional clients, designed to help **exchanges and regulators** identify suspicious accounts. However, **ordinary on-chain users** face compliance risks that are often passive, with no tools to address them.

**Typical scenarios:**

Scenario 1: Incoming Transfer Taint
Blacklisted address A → transfers to → ordinary user B → transfers to → centralized exchange
Result: The exchange's AML system detects that B's deposit originates from high-risk address A; B's account is frozen or flagged as high-risk, despite B's complete unawareness.

Scenario 2: Cross-Chain Taint Propagation
Blacklisted address A → mixer → cross-chain bridge → relay address C → ordinary user B
Result: B does not know that the upstream of the transfer from C contains a blacklisted address. The exchange's 2-hop or 3-hop scan may trigger a risk alert.

This phenomenon is known as **"Passive Taint"** or **"Innocent Third-Party Harm."** The core contradiction is:

- Blockchain is publicly transparent — anyone can theoretically verify an address's history — but **analytical capabilities are monopolized by commercial companies**, leaving individual users unable to perform equivalent risk self-checks before initiating transactions.
- Exchange compliance policies are **opaque and inconsistent**: some trace 2 hops, others trace 5 hops; users have no idea how deep their deposits will be scrutinized.
- Launderers intentionally **distribute funds to numerous ordinary addresses** (the Money Mule pattern described above), unknowingly making these addresses part of the laundering chain, which then deposit into exchanges to bypass direct blacklist detection.

BlockSec's analysis found that **54% of blacklisted addresses had completed asset transfers before being frozen** — meaning substantial "tainted funds" are already circulating among normal user addresses, while these users remain completely unaware.

This research aims to fill this tool gap: providing **ordinary on-chain users** with a self-service address risk assessment tool, enabling them to determine whether a counterparty address has fund associations with known violations before initiating transactions or accepting transfers, thereby proactively avoiding passive taint risk.

### 2.4 Formal Definition of Research Questions

Synthesizing the above background, this research addresses the following core question:

**Given a target blockchain address, how can we systematically identify whether it has fund associations with known blacklisted addresses, and what is the risk level and confidence of such associations — particularly when fund paths cross multiple blockchains, pass through mixers, or involve multi-layer relays?**

Specifically, this research addresses four sub-questions:

- **Sub-question 1 (Classification)**: How to distinguish between "genuinely untraceable concealment behaviors" (mixers, opaque bridges) and "penetratable cross-chain behaviors" (transparent bridges)?
- **Sub-question 2 (Depth)**: In multi-hop tracing, how to find a reasonable balance between "insufficient tracing depth (evadable)" and "excessive tracing overhead (false positives on normal users)"?
- **Sub-question 3 (Scoring)**: How to quantify multi-hop, multi-chain tracing results into interpretable risk scores that enable horizontal comparison across different addresses?
- **Sub-question 4 (Proportion)**: When an address holds both legitimate and tainted funds simultaneously, how to distinguish the proportion of "dirty money" from "clean money," avoiding blanket high-risk classification of the entire address?

---

## 3. Related Literature Review

### 3.1 Blockchain Transaction Graphs and AML Detection

Modeling blockchain transactions as graphs is the mainstream paradigm in current AML research. The most influential foundational work in this direction comes from Weber et al. (2019), who released the **Elliptic Dataset** — a Bitcoin transaction graph containing 203,769 nodes and 234,355 directed edges, with approximately 4,500 nodes labeled as "illicit," first proposing the application of Graph Convolutional Networks (GCN) to AML classification tasks at the KDD 2019 Anomaly Detection in Finance Workshop [1]. This dataset remains the most widely used AML benchmark in academia. In 2024, the same team released **Elliptic2**, providing community-level annotations to support morphological analysis of entire laundering subgraphs (rather than individual transactions) [2].

However, these studies share a common limitation: **all are based on single chains (primarily Bitcoin or Ethereum), without discussing how to continue tracing after funds cross chains.** Additionally, address clustering is a core capability of commercial tools (Chainalysis, Elliptic), aiming to consolidate multiple addresses controlled by the same entity, but the relevant algorithmic details are not public, and academic open-source implementations primarily rely on heuristic methods with limited coverage under Ethereum's account model. This research does not directly implement address clustering but uses interaction frequency as a proxy metric during child node generation to partially compensate for this gap.

### 3.2 Cross-Chain Transaction Tracing and Traceability

This direction provides the most direct academic context for this research and is one of the fastest-growing research areas in recent years.

#### 3.2.1 Transparent Bridge Tracing Methods

**Mazorra et al. (2023/2024)** in *Tracing Cross-chain Transactions between EVM-based Blockchains: An Analysis of Ethereum-Polygon Bridges* (published in the Ledger journal) proposed a set of heuristic matching algorithms for cross-chain transactions between EVM-compatible chains [3]. The core insight is: **between EVM-compatible chains, user addresses remain consistent across different chains** (e.g., the same `0xABCD...` address on Ethereum and Polygon is controlled by the same key), enabling matching of Lock events on the source chain with Mint/Release events on the target chain through a "time window + amount + token type" combination algorithm. This research achieved a **99.65%** deposit matching rate and **92.78%** withdrawal matching rate on over 2 million cross-chain transactions spanning August 2020 to August 2023.

**Sun et al. (2025)** in *Track and Trace: Automatically Uncovering Cross-chain Transactions in the Multi-blockchain Ecosystems* (arXiv 2504.01822) systematically analyzed **12 major cross-chain bridges** (including Stargate, Celer cBridge, Wormhole, Synapse, etc.), covering Ethereum source chain data from April 2021 to March 2024, and proposed a general framework for automatically identifying cross-chain transactions [4]. A key finding was that transparency varies dramatically across bridges — message-passing protocol-based bridges (e.g., LayerZero) can obtain counterpart transaction hashes directly via API, while liquidity pool-based bridges (e.g., Synapse) make it nearly impossible to find explicit input-output correspondences in on-chain data.

**A Survey of Transaction Tracing Techniques for Blockchain Systems** (arXiv 2510.09624) provides a more macro-level overview of blockchain transaction tracing techniques, categorizing existing methods into: on-chain event correlation, API-assisted tracing, statistical inference, and machine learning, noting that cross-chain tracing is currently the area most lacking systematic research [5].

#### 3.2.2 Transparent vs. Opaque Bridges: The Importance of Traceability Classification

This research argues that distinguishing **transparent bridges (traceable)** from **opaque bridges (opaque)** is a prerequisite for effective fund tracing, but this classification has not been adequately discussed in existing academic literature. The following comparison illustrates the fundamental difference (table synthesized from [3], [4], [5]):

| Dimension | Transparent Bridge | Opaque Bridge |
|-----------|-------------------|---------------|
| **Mechanism** | Message-passing protocols (LayerZero, Wormhole) or Rollup official bridges | Liquidity pools (Synapse), market-maker models (Orbiter, Owlto) |
| **On-chain correspondence** | Unique transferId or shared txHash links both ends | User funds enter a shared pool; market makers independently transfer on the target chain; no linkage possible |
| **Analogy** | Bank wire transfer (with reference number) | Cash deposited at ATM, withdrawn by a different person (no linkage) |
| **Laundering risk** | Low (fund flows can be traced and audited) | High (equivalent to a mixer; fund flows are untraceable) |
| **Typical examples** | Stargate, Hop Protocol, Arbitrum/Optimism official bridges | Multichain (collapsed), Orbiter Finance, Synapse |

Multichain (formerly one of the largest cross-chain bridges) collapsed in 2023 due to internal issues, resulting in approximately **$127 million** in lost assets — an event that itself exposed the fundamental deficiencies of opaque bridges in transparency and auditability.

#### 3.2.3 Cross-Chain Laundering Paths in Criminal Cases

The laundering paths used by Lazarus Group in the Ronin Bridge and Harmony Bridge attacks (detailed in Section 1.2) are the most studied cross-chain laundering cases to date. Chainalysis's on-chain tracing showed that after multiple layers of bridging and mixing operations, most funds ultimately flowed to OTC brokers or high-risk exchanges; law enforcement eventually recovered approximately **$30 million** — only about 4.8% of the total stolen from Ronin.

This figure reveals the current ceiling of tracing capabilities: even the most well-resourced commercial tools achieve extremely low recovery rates when facing the combination of mixers and multi-chain hops. This is precisely why improving automated tracing tools has significant academic research value.

### 3.3 Multi-Hop Risk Scoring and Tree-Based Tracing: The Most Directly Related Prior Work

This section reviews three works most directly related to the technical approach of this research, with a detailed comparison table at the end.

#### 3.3.1 Foundational Work: Bitcoin Transaction Risk Scoring (2014)

**Möser, Böhme & Breuker (2014)** in *Towards Risk Scoring of Bitcoin Transactions* (Financial Cryptography 2014) first formalized "propagating risk from known illicit addresses along the transaction graph" as a research problem [6]. The work proposed two propagation strategies:

- **Poison (full taint)**: If any illicit input exists in the fund sources, the output is considered fully tainted
- **Haircut (proportional)**: Taint proportion equals the ratio of illicit source funds to total inputs, gradually diluted through mixing

This is the earliest academic work discussing "multi-hop risk" and has been cited by virtually all subsequent Taint Analysis research. Its limitation lies in its 2014 publication date, with the research scope limited to Bitcoin single-chain, without addressing cross-chain bridges or mixer-specific processing logic.

#### 3.3.2 TaintRank: PageRank-Style Taint Propagation (2019)

**Hercog & Povšea (2019)** in *Taint analysis of the Bitcoin network* (arXiv:1907.01538) proposed the **TaintRank** algorithm [7], analogizing taint propagation to PageRank:

- Construct a weighted directed graph with addresses as nodes and transactions as directed edges
- Each node's taint value is the weighted sum of all upstream nodes' taint values
- Taint naturally decays with propagation distance, with the final distribution following a power-law shape

TaintRank scores the entire Bitcoin network in a **batch, global** manner, producing a 0-1 taint index for each address. Unlike this research, it is an offline batch-processing algorithm that does not support real-time tree-based queries rooted at individual addresses, nor cross-chain scenarios.

#### 3.3.3 Transaction Proximity: Circle's BFS Application on the Ethereum Full Graph (2025)

**Liao, Zeng, Belenkiy & Hirshman (2025)** from USDC issuer Circle, in *Transaction Proximity: A Graph-Based Approach to Blockchain Fraud Prevention* (arXiv:2505.24284), applied the BFS approach to the entire Ethereum historical graph [8]:

- Data scale: **206 million nodes, 442 million edges**, covering all Ethereum transactions from genesis to May 2024
- BFS depth limit: **5 hops** (covering 98.2% of active USDC holders)
- Core metrics: **Transaction Proximity** (shortest hop count to a regulated exchange) and **EAI (Easily Attainable Identities)** (addresses directly connected to exchanges)
- Key finding: 83% of known attacker addresses are not EAI; 21% are more than 5 hops from any regulated exchange — demonstrating that criminal addresses tend to be structurally distant from "normal circulation nodes"

Notably, this paper's risk perspective is **opposite but complementary** to this research: Transaction Proximity measures "distance from legitimate anchors" (closer = more legitimate), while this research measures "distance from illicit anchors" (closer = more dangerous). The two methods can theoretically be fused to provide confidence from both directions for the same address.

#### 3.3.4 Systematic Comparison with This Research

The following table compares this research with existing approaches (table compiled from [6], [7], [8]):

| Feature | Möser 2014 | TaintRank 2019 | Tx Proximity 2025 | Chainalysis (Industry) | **This Research** |
|---------|:---:|:---:|:---:|:---:|:---:|
| Cross-chain tracing | No | No | No | Partial (undisclosed) | Yes |
| Transparent vs. opaque bridge classification | No | No | No | No | Yes |
| Adaptive depth (deepen suspicious branches) | No | No | No (fixed 5 hops) | No | Yes |
| Real-time single-address query | No (paper) | No (batch) | No (offline) | Yes (commercial) | Yes |
| Targeting ordinary users (non-institutional) | — | — | — | No | Yes |
| Blacklist anchor: USDT freeze list | No | No | No | Partial | Yes |
| Explicit hop decay coefficient | Yes | Implicit | No | Undisclosed | Yes (×0.6) |
| Mixer termination logic | No | No | No | Yes | Yes |
| Tree visualization + Mermaid export | No | No | No | Yes (commercial) | Yes |
| Open-source reproducible | — | Partial | Partial | No | Yes |

The comparison shows that this research's core incremental contribution lies in its **cross-chain tracing framework** (including bridge traceability classification) and **adaptive depth** mechanism, both of which have no direct precedent in existing academic literature. Transaction Proximity (2025) is the methodologically closest work, but it covers neither cross-chain capability nor the ordinary-user-facing positioning.

### 3.4 Regulatory Frameworks and Real-World Needs

The **Financial Action Task Force (FATF)** incorporated virtual assets (VA) and virtual asset service providers (VASPs) into its AML/CFT standards framework (Recommendation 15) in 2019. In its 2023 targeted update report, FATF noted that among 151 member jurisdictions, **more than half had not yet implemented the "Travel Rule,"** and 75% were partially or non-compliant with R.15 [9]. The FATF 2023 report also specifically highlighted the notable growth in stablecoin use by DPRK actors, terrorism financiers, and drug traffickers.

This background demonstrates that significant gaps remain in the current compliance system, and automated, interpretable on-chain tracing tools have clear real-world demand beyond academic research interest.

Complementing the FATF framework, the **EU AI Act (Regulation 2024/1689)**, which entered into force in 2024, classifies AI systems used for AML/CFT compliance as **"High-Risk AI Systems" (Annex III)** [18]. The Act requires such systems to satisfy: (1) sufficient transparency of decision logic, (2) human oversight mechanisms, and (3) ability to provide complete decision explanations to regulators and end users. Violators face fines up to **EUR 35 million or 7% of global turnover**. This means "black box" models (e.g., pure GNN classifiers) without an explanation layer will increasingly struggle to meet European market compliance requirements — providing clear institutional motivation for interpretable AML systems.

### 3.5 Interpretability of AML Systems and the LLM Explanation Layer

#### 3.5.1 Interpretability as a Regulatory Mandate

The interpretability of AML detection systems is not "nice to have" but a hard regulatory requirement. FATF Recommendation 20 requires financial institutions to include a textual explanation of "why the transaction is considered suspicious" when filing Suspicious Transaction Reports (STRs). The US Bank Secrecy Act similarly requires Suspicious Activity Reports (SARs) to describe suspicious behavioral patterns in natural language. The aforementioned EU AI Act further extends this explanation obligation to AI systems themselves [18].

This creates a core tension: **the models with the best detection performance (GNN, deep learning) are precisely the least interpretable; while the most interpretable methods (rule engines) have limited detection capability.**

#### 3.5.2 XAI (Explainable AI) in AML

Research on integrating explainability tools (SHAP, LIME, etc.) into AML models has grown rapidly in recent years. Kute et al. (2026) proposed an end-to-end reproducible SHAP framework for explaining AML model decisions, emphasizing the integration of fairness and regulatory compliance [17]. This work systematically demonstrated how to transform feature importance attributions into reports understandable by compliance departments.

In the blockchain context, Watson, Richards & Schiff (2025) proposed a representative three-layer architecture (architecture from [14]):

Layer 1: GNN detection layer (GCN-GRU hybrid model, accuracy 0.9470, AUC-ROC 0.9807)
Layer 2: GraphLIME attribution layer (identifies which features drove the classification result)
Layer 3: LLM explanation layer (converts attribution results into natural language narratives)

This architecture was validated on the Elliptic++ dataset, in a paper titled *Explain First, Trust Later: LLM-Augmented Explanations for Graph-Based Crypto Anomaly Detection*, representing the state-of-the-art paradigm of GNN + XAI + LLM three-layer fusion.

#### 3.5.3 Standalone LLM Applications in On-Chain Transaction Explanation

Beyond serving as an explanation layer for GNNs, LLMs have also been directly applied to on-chain transaction data analysis and explanation. Nicholls et al. (2024) in *Large Language Model XAI Approach for Illicit Activity Investigation in Bitcoin* (published in Springer *Neural Computing and Applications*, IF 4.7) [15] demonstrated a method independent of traditional ML detectors: directly feeding Bitcoin transaction data to LLMs to generate natural language narratives, then extracting narrative embeddings to compute similarity, thereby discovering other illicit transactions. The significance of this approach lies in proving that LLMs can produce "useful, actionable explanations" from on-chain transaction data, and that these explanations themselves can serve as features for downstream detection.

Sun et al. (2024) published the first systematic review of LLM applications in blockchain security, *Large Language Models for Blockchain Security: A Systematic Literature Review* [16], covering anomaly detection, smart contract auditing, transaction analysis, and other directions, providing a structured overview of the field's research landscape.

#### 3.5.4 Relationship to This Research

This research's rule engine (BFS tracing + risk scoring) has inherent advantages in interpretability: every decision step (why a node was flagged as suspect, what the risk path is, how decay was calculated) has a complete causal chain, with no "black box" problem. This makes introducing LLMs as an explanation layer more straightforward — no need for post-hoc attribution tools like SHAP/LIME to "reverse-engineer" model decisions; instead, directly converting existing structured tracing results (JSON) into natural language risk reports, inherently meeting the narrative requirements of SARs/STRs.

### 3.6 Privacy-Preserving Compliance: Zero-Knowledge Proofs and Privacy Pools

#### 3.6.1 The Core Tension Between Privacy and Compliance

A fundamental tension exists in blockchain AML: **privacy protection and compliance auditing have long been viewed as mutually exclusive poles.** Users have legitimate privacy needs (not exposing asset positions and transaction details), but compliance frameworks require fund source traceability.

In August 2022, US OFAC sanctioned Tornado Cash smart contracts for the first time — the first-ever sanction imposed on **open-source, immutable code** rather than a person or organization, sparking widespread academic and legal discussion. Brownworth, Durfee, Lee & Martin (2024) provided systematic empirical analysis in a New York Federal Reserve working paper [21]: transaction volumes and user diversity dropped immediately after the sanction announcement, but net inflows recovered and even exceeded pre-sanction levels within months; the number of block validators processing Tornado Cash transactions continued to shrink, indicating that **censorship resistance is fragile**. In November 2024, the US Fifth Circuit Court of Appeals ruled that OFAC's sanction of immutable smart contracts exceeded its statutory authority; in March 2025, the US Treasury officially lifted the sanctions on Tornado Cash.

These events demonstrate that **simply sanctioning privacy tools is not a sustainable solution**, and that technical approaches to balancing privacy and compliance are needed.

#### 3.6.2 Privacy Pools: A Practical Equilibrium Between Privacy and Compliance

**Buterin, Illum, Nadler, Schär & Soleimani (2023)** in *Blockchain Privacy and Regulatory Compliance: Towards a Practical Equilibrium* (published in *Blockchain: Research and Applications*) [19] proposed the **Privacy Pools** protocol and the core concept of **Association Sets**:

- **Deposit phase**: Users deposit funds into the privacy pool (similar to Tornado Cash)
- **Withdrawal phase**: Users select an Association Set and use a zero-knowledge proof (ZKP) to prove "my deposit belongs to this set" without revealing which specific deposit
- **Association Set semantics**:
  - **Inclusion set**: "My deposit belongs to {the set of deposits from known legitimate sources}"
  - **Exclusion set**: "My deposit does not belong to {the set of deposits from OFAC-sanctioned addresses}"

This protocol first demonstrated academically that **"privacy" and "compliance" need not be mutually exclusive** — users can protect transaction details while proving the legitimacy of their fund sources to verifiers. Privacy Pools v1 launched on Ethereum mainnet in 2024.

#### 3.6.3 Practice and Limitations of Proof of Innocence

The core idea of Privacy Pools gave rise to **Proof of Innocence** in engineering practice. The Railgun protocol was first to deploy **Private Proofs of Innocence (PPOI)**: when users deposit tokens (shield), the wallet automatically generates a ZK proof that the token "does not belong to a predefined list of illicit transactions/addresses." This proof is verified by decentralized POI nodes, with the entire process end-to-end encrypted, exposing neither the user's 0zk address, balance, nor transaction history.

However, Constantinides & Cartlidge (2025) in *zkMixer: A Configurable Zero-Knowledge Mixer with Anti-Money Laundering Consensus Protocols* (accepted at IEEE DAPPS 2025) [20] identified fundamental limitations of Proof of Innocence in practice: **PoI depends on the completeness and timeliness of the blacklist** — if a deposit is flagged as illicit after passing the PoI check, the deposit has already entered the privacy pool and is irrevocable. The paper proposed an alternative: using a consensus mechanism to collectively verify deposits **before** they enter the mixing pool; if verification fails, the deposit can be frozen or returned.

#### 3.6.4 Cryptographic Techniques for Privacy-Preserving AML

Beyond ZKP, other cryptographic methods have been explored for privacy-preserving AML:

- **Fully Homomorphic Encryption (FHE)**: Effendi & Chattopadhyay (2024) in *Privacy-Preserving Graph-Based Machine Learning with Fully Homomorphic Encryption for Collaborative Anti-Money Laundering* (SPACE 2024) [23] demonstrated using FHE to execute graph machine learning directly on encrypted data (XGBoost achieving 99%+ accuracy), enabling multiple financial institutions to collaboratively perform AML detection without sharing raw data.
- **ZKP Middleware**: Chaudhary (2023) proposed the zkFi framework [22], encapsulating zero-knowledge proofs as plug-and-play compliance plugins for DeFi protocols, lowering the barrier to ZKP integration.

#### 3.6.5 Relationship to This Research

This research currently performs risk tracing based on public on-chain data, with tracing results (complete BFS path trees) exposing users' transaction association information. From a privacy perspective, a natural evolutionary direction is: **users run tracing locally and generate a ZK proof — "this address has no blacklist association within N hops" — without revealing the specific path.** The verifier (exchange or counterparty) sees only whether the proof is valid, not any intermediate addresses involved in the tracing process.

The main challenges facing this direction include: (1) real-time blacklist update problem — proof validity depends on a specific blacklist snapshot; (2) computational overhead — circuit complexity for generating ZK proofs over BFS trace trees is high; (3) Association Set builder trust problem — who decides which addresses belong to the "legitimate" set. These are open academic questions beyond the current scope of this research but have clear value as long-term research directions.

---

## 4. Research Methodology Overview

Based on the above background and literature, this research proposes a **Multi-Chain Risk Trace Graph** rooted at individual addresses, with the following core design decisions:

### 4.1 Bridge Traceability Classification as the Foundation of the Tracing Framework

Unlike existing research, this study treats cross-chain bridge traceability as a first-class citizen in the analytical framework, classifying bridges as:
- **Transparent bridges**: Obtain counterpart addresses through protocol APIs (LayerZero Scan, Hop, Wormhole, etc.) or event logs (Rollup bridges), continuing analysis on the target chain
- **Opaque bridges**: Treated as equivalent to mixers, flagged as high-risk "tracing breakpoint" terminal nodes

This classification directly addresses the traceability disparity issues raised by Sun et al. (2025) in their multi-bridge study, providing an operationally implementable solution at the engineering level.

### 4.2 Tree-Based Tracing with Adaptive Depth

This research uses BFS (Breadth-First Search) to construct fund association trees, introducing an **Adaptive Depth** mechanism:
- Normal branches use the standard maximum depth (default 3 hops)
- Branches where suspicious indicators are found (blacklist contact / mixer usage / opaque bridge usage) receive a `depth_bonus` of additional hops (default +1)

This design provides a data-driven balance point between "insufficient tracing" and "excessive tracing," rather than an arbitrary fixed threshold.

### 4.3 Hop-Decay Risk Scoring

Following the AML practice consensus that "direct association risk is higher than indirect association risk," this research introduces a **0.6× per-hop** risk decay coefficient:

| Association Distance | Effective Risk After Decay (from 100) | Risk Level |
|---------------------|---------------------------------------|------------|
| Direct contact (1 hop) | 60 | HIGH |
| 2 hops | 36 | MEDIUM |
| 3 hops | 21.6 | LOW |
| 4+ hops | ≤ 13 | Reference |

### 4.4 Relay Address Identification

Targeting "clean address relay" (Money Mule) — a common laundering technique — the system performs a secondary scan of the entire tree after BFS expansion: if an ostensibly clean address's subtree contains blacklist hits, it is flagged as a "suspect relay address," and a subtree contamination score is calculated for that address.

---

## 5. Current Progress

### 5.1 Completed Work

The research has completed implementation of the following core modules:

**`aml_analyzer.py` — Single-Address Analysis Engine**
- Integrated Etherscan API and TronScan API, fully supporting Ethereum and Tron chains
- Implemented USDT blacklist detection (8,500+ addresses, covering Ethereum and Tron)
- Implemented cross-chain bridge registry (`BRIDGE_REGISTRY`), containing 20+ bridge contracts, distinguishing transparent/opaque
- Implemented mixer identification (Tornado Cash and 10 other contracts)
- Implemented `BridgeTracer`: obtains cross-chain counterpart addresses via LayerZero Scan API
- Implemented risk scoring (0-100 scale)

**`trace_graph.py` — Recursive Trace Graph Engine**
- BFS-based multi-hop fund association tree construction
- Transparent bridge → chain switching and continued tracing
- Adaptive depth (`depth_bonus`)
- Subtree risk propagation (0.6 decay per hop)
- "Suspect relay address" secondary flagging
- JSON export and Mermaid visualization graph export

### 5.2 Test Data Selection

The system is testable. The following are representative addresses planned for testing:

| Address | Type | Expected Result |
|---------|------|-----------------|
| `0x098b716b8aaf21512996dc57eb0615e2383e2f96` | Ronin Bridge attacker (Lazarus Group), on USDT blacklist | Direct blacklist hit, tree terminates at first layer |
| `0x7f367cc41522ce07553e823bf3be79a889debe1b` | Lazarus Group-associated address, on USDT blacklist | Blacklist hit |
| Upstream address with direct interaction with above | One-hop association (via Etherscan query) | Risk score ≈ 60, flagged as suspect relay |
| Test address using Stargate Finance cross-chain | Transparent bridge user | Chain switch to target chain for continued analysis |

### 5.3 Future Plans

**Near-Term Engineering Improvements (Short-Term):**
- Supplement OFAC SDN list (covering addresses from Harmony attack not frozen by Tether)
- Implement Hop Protocol and Across Protocol counterpart address resolution
- Conduct end-to-end tracing tests on known laundering cases (Ronin/Harmony) to verify whether the tree structure can reconstruct actual laundering paths
- Introduce time window filtering to avoid false positives from ancient transactions

**Medium-Term Research Direction: Taint Proportion Analysis**

The current system applies holistic risk flagging to addresses that "have used a mixer" or "received blacklisted transfers," without distinguishing the proportion of tainted versus legitimate funds. In practice, this causes false positives: an address that received 1 USDT from a blacklisted transfer but also holds 10,000 USDT in entirely legitimate funds should not be treated equivalently to a direct laundering address.

**Taint Analysis** is an established research direction in blockchain forensics, with the core question being: given an address, what proportion of its assets can be traced to known illicit sources? Typical methods include:
- **FIFO method**: First-in-first-out, assuming the earliest received funds are spent first
- **Haircut method**: Taint proportion is proportionally diluted with each mixing event
- **Poison method**: Any illicit source taints the entire output (most conservative)

Introducing Taint Analysis into this system would upgrade the current binary judgment (risk/non-risk) to **proportional confidence scoring**, better aligned with practical compliance needs and the "distinguish dirty money from clean money" objective raised in meetings.

**Long-Term Research Directions: LLM Explanation Layer and Compliance Privacy**

- **LLM Risk Explanation Layer**: Feed the existing system's JSON tracing results into an LLM to automatically generate natural language risk reports meeting SAR/STR regulatory requirements (see Section 3.5 literature analysis). This direction has been validated by Watson et al. [14] and Nicholls et al. [15] on Elliptic++ and Bitcoin datasets.

- **ZKP Compliance Proof**: After running tracing locally, users generate a zero-knowledge proof — "this address has no blacklist association within N hops" — without revealing the specific path. The theoretical foundation comes from Buterin et al.'s Privacy Pools [19], with practical limitations and alternatives systematically analyzed by Constantinides & Cartlidge [20].

These two directions correspond to the EU AI Act's interpretability requirements [18] and the long-term need for privacy-preserving compliance, respectively. They exceed the current midterm implementation scope but are worth proposing as research topics in this report.

---

## 6. System Improvements: From Rule Engine to Machine Learning

### 6.1 Systematic Defects in the Existing Rule Engine

After code review of `aml_analyzer.py` and `trace_graph.py`, five systematic issues affecting analysis accuracy were identified, falling into two categories:

**Category 1: Sampling Bias**

| ID | Issue | Impact | Fix |
|:---:|-------|--------|-----|
| P1 | `MAX_TX_FETCH=100`, only fetching latest 100 transactions | Active addresses' historical dirty transactions truncated; launderers can use "dilution attacks" to push dirty transactions out of the window with junk transactions | Increased to 500 with truncation warning |
| P2 | `txlist + tokentx` direct concatenation causes double-counting of the same transaction | Counterparty interaction frequency inflated, ranking distorted | Deduplication by `(hash, from, to)` tuple |
| P5 | Counterparty ranking purely by interaction frequency | DEX Routers and exchange hot wallets fill all slots, burying low-frequency but high-value suspicious addresses | Amount-weighted composite scoring, excluding known DEX addresses |

**Category 2: Graph Distortion**

| ID | Issue | Impact | Fix |
|:---:|-------|--------|-----|
| P3 | Bridge/mixer detection only checks the `to` field | Extracting funds from Tornado Cash (`from=mixer`) is completely undetected | Bidirectional `from` and `to` detection, recording direction (IN/OUT) |
| P4 | `visited` set silently discards convergence paths | Multiple paths pointing to the same node (scatter→converge laundering pattern) become invisible | Preserve convergence information (`converge_from`, `in_degree`), record without re-expansion |

**Core insight: The common effect of these five issues is that launderers' adversarial strategies (creating noise, extracting from mixers, scatter-then-converge) precisely exploit the system's blind spots.** The sampling strategy and launderers' adversarial strategies move in opposite directions — launderers create noise to dilute signals, while the system prioritizes noise and filters out signals.

### 6.2 Machine Learning Workflow

#### 6.2.1 Motivation and Method Selection

Following the **StableAML** paper by Juvinski & Li (2026) [24], which trained tree ensemble models on 16,433 labeled addresses using 68 behavioral features (CatBoost achieving Macro-F1 = 0.9775), a key finding was: **domain-specific feature engineering matters more than complex graph algorithms** — tree ensembles (CatBoost, F1=0.9775) significantly outperformed graph neural networks (GraphSAGE, F1=0.8048), because stablecoin transaction graphs are extremely sparse (density < 0.01), rendering GNN message passing ineffective at propagating information.

Based on this finding, this research adopts a **feature engineering + tree ensemble model** approach rather than GNNs.

#### 6.2.2 Data Collection

Data was sourced from two channels:

- **Blocklisted class** (100 addresses): Ethereum addresses from the project's `usdt_blacklist.csv` (Tether's official freeze list, ~8,500 entries)
- **Normal class** (50 addresses): Randomly sampled active addresses from recent Ethereum block USDT Transfer events, excluding known contracts (bridges, mixers, exchanges, zero addresses) and blacklisted addresses
- **Sanctioned class** (10 addresses): Automatically downloaded from the OFAC SDN sanctions list, extracting Ethereum addresses

For each address, the Etherscan getLogs API was used to fetch all USDT/USDC **Transfer events** (ERC-20 standard events), querying bidirectionally for `sent` (topic[1]=address) and `received` (topic[2]=address). The choice of getLogs over txlist was motivated by: (1) avoiding txlist/tokentx cross-duplication (P2); (2) Transfer events are the sole authoritative record of token movements; (3) native bidirectional query support (addressing P3).

#### 6.2.3 Feature Engineering

Following StableAML's four-category feature framework (feature framework adapted from [24]), **61 behavioral features** were extracted from raw Transfer events:

| Category | Count | Representative Features | Data Source |
|----------|:-----:|------------------------|-------------|
| Interaction Features | 18 | `sent_to_mixer`, `received_from_mixer`, `has_flagged_interaction` | Matched against project's existing address registries (BRIDGE_REGISTRY, MIXER_CONTRACTS, etc.) |
| Transfer Features | 19 | `transfers_over_10k`, `drain_ratio`, `repeated_amount_ratio` | Pure amount/count statistics |
| Network Features | 10 | `in_degree`, `out_degree`, `counterparty_flagged_ratio`, `has_proxy_behavior` | Computed from from/to sets |
| Temporal Features | 8 | `has_daily_burst`, `rapid_tx_ratio`, `hour_concentration` | Computed after timestamp sorting |

`has_proxy_behavior` detects "received then transferred the same amount (±5%) within 24 hours" patterns (peeling chain relay signature); `repeated_amount_ratio` detects repeated-amount transfers (peeling chain signal).

#### 6.2.4 Model Training and Evaluation

Using 5-Fold Stratified Cross Validation, four tree ensemble models were compared:

| Model | Macro-F1 | PR-AUC |
|-------|:--------:|:------:|
| **RandomForest** | **0.919** | **0.949** |
| XGBoost | 0.886 | 0.917 |
| CatBoost | 0.872 | 0.937 |
| LightGBM | 0.865 | 0.932 |

Confusion matrix for the best model (RandomForest):

|  | Predicted blocklisted | Predicted normal |
|--|:-:|:-:|
| Actual blocklisted | 93 | 7 |
| Actual normal | 4 | 46 |

Cross-model consensus Top 5 important features:

- **`drain_ratio`** (balance drain rate) — blocklisted mean 0.21 vs normal 1.37 (frozen addresses cannot transfer funds out)
- **`total_sent_amount`** (total sent amount) — large-scale fund movement is the core signal
- **`counterparty_flagged_ratio`** (counterparty flagged ratio) — reflects the value of KYC/label data
- **`out_degree`** (out-degree) — blacklisted addresses have significantly lower out-degree than normal addresses
- **`in_out_ratio`** (inflow/outflow ratio) — fund flow symmetry

#### 6.2.5 Model Integration

The trained model is integrated into the existing system through the `MLRiskScorer` class:

Original risk score = Pure rule engine (hardcoded weights)
New risk score = Rule engine × 0.4 + ML model predict_proba × 0.6

Design rationale for the hybrid strategy:
- Rule engine retains 40% weight: Ensures known high-risk signals (mixers, direct blacklist association) are not underestimated by the ML model
- ML model takes 60% weight: Provides supplementary signals for behavioral patterns not covered by the rule engine (temporal anomalies, amount distributions, network topology)
- Graceful degradation: Automatically falls back to pure rule engine if model files do not exist

### 6.3 Limitations

- **Limited data**: 150 samples vs StableAML's 16,433. Expanding the dataset is the most direct means to improve model generalization.
- **Unverified Normal class**: Randomly sampled "normal" addresses from on-chain may include unlabeled laundering addresses (label noise).
- **USDT/USDC Transfer events only**: Launderers who swap to ETH or other tokens escape the analysis scope.
- **Real-time feature extraction**: The current ML scorer extracts features from `RiskReport` as a subset of existing report information; some features (e.g., temporal features) require complete Transfer event data for precise calculation.
- **Binary classification limitation**: Currently only distinguishes blocklisted/normal, without sanctioned, cybercrime, or other sub-categories.

---

## References

### AML Graph Analysis Foundations

[1] Weber, M., Domeniconi, G., Chen, J., Weidele, D. K. I., Bellei, C., Robinson, T., & Leiserson, C. E. (2019). Anti-Money Laundering in Bitcoin: Experimenting with Graph Convolutional Networks for Financial Forensics. KDD Workshop on Anomaly Detection in Finance. arXiv:1908.02591.

[2] Bellei, C. et al. (2024). The Shape of Money Laundering: Subgraph Representation Learning on the Blockchain with the Elliptic2 Dataset. arXiv:2404.19109.

### Cross-Chain Tracing

[3] Mazorra, B. et al. (2023). Tracing Cross-Chain Transactions Between EVM-Based Blockchains: An Analysis of Ethereum-Polygon Bridges. Ledger Journal. arXiv:2504.15449.

[4] Sun, X. et al. (2025). Track and Trace: Automatically Uncovering Cross-chain Transactions in the Multi-blockchain Ecosystems. arXiv:2504.01822.

[5] Ren, J. et al. (2025). A Survey of Transaction Tracing Techniques for Blockchain Systems. arXiv:2510.09624.

### Multi-Hop Risk Scoring

[6] Möser, M., Böhme, R., & Breuker, D. (2014). Towards Risk Scoring of Bitcoin Transactions. Financial Cryptography and Data Security Workshops (FC 2014). Springer.

[7] Hercog, U., & Povšea, A. (2019). Taint Analysis of the Bitcoin Network. arXiv:1907.01538.

[8] Liao, G., Zeng, Z., Belenkiy, M., & Hirshman, J. (2025). Transaction Proximity: A Graph-Based Approach to Blockchain Fraud Prevention. Circle Research. arXiv:2505.24284.

### Industry Reports and Data Sources

[9] FATF (2023). Targeted Update on Implementation of the FATF Standards on Virtual Assets and Virtual Asset Service Providers. Financial Action Task Force. June 2023.

[10] Chainalysis (2024). 2024 Crypto Money Laundering Report. Chainalysis Inc.

[11] Elliptic (2023). $7 Billion in Crypto Laundered Through Cross-Chain Services. Elliptic Enterprise Ltd.

[12] BlockSec (2023). Following the Frozen: An On-Chain Analysis of USDT Blacklisting and Its Links to Terrorist Financing. BlockSec Blog.

[13] Weber, M. & Bellei, C. (2019). Elliptic Data Set. Kaggle / Elliptic.

### LLM Explanation Layer and Explainable AML

[14] Watson, A., Richards, G., & Schiff, D. (2025). Explain First, Trust Later: LLM-Augmented Explanations for Graph-Based Crypto Anomaly Detection. arXiv:2506.14933.

[15] Nicholls, J. et al. (2024). Large Language Model XAI Approach for Illicit Activity Investigation in Bitcoin. Neural Computing and Applications. Springer.

[16] Sun, H. et al. (2024). Large Language Models for Blockchain Security: A Systematic Literature Review. arXiv:2403.14280.

[17] Kute, D. et al. (2026). Explainable and Fair Anti-Money Laundering Models Using a Reproducible SHAP Framework for Financial Institutions. Discover Artificial Intelligence. Springer.

### Regulatory Frameworks

[18] European Union (2024). Regulation (EU) 2024/1689 — AI Act, Annex III: High-Risk AI Systems.

### ZKP Compliance Proofs and Privacy-Preserving AML

[19] Buterin, V., Illum, J., Nadler, M., Schär, F., & Soleimani, A. (2023). Blockchain Privacy and Regulatory Compliance: Towards a Practical Equilibrium. Blockchain: Research and Applications, 5(1), 100176.

[20] Constantinides, T. & Cartlidge, J. (2025). zkMixer: A Configurable Zero-Knowledge Mixer with Anti-Money Laundering Consensus Protocols. arXiv:2503.14729. Accepted at IEEE DAPPS 2025.

[21] Brownworth, A., Durfee, J., Lee, M., & Martin, A. (2024). Regulating Decentralized Systems: Evidence from Sanctions on Tornado Cash. Federal Reserve Bank of New York Staff Reports, No. 1112.

[22] Chaudhary, A. (2023). zkFi: Privacy-Preserving and Regulation Compliant Transactions using Zero Knowledge Proofs. arXiv:2307.00521.

[23] Effendi, F. & Chattopadhyay, A. (2024). Privacy-Preserving Graph-Based Machine Learning with Fully Homomorphic Encryption for Collaborative Anti-Money Laundering. SPACE 2024. arXiv:2411.02926.

[24] Juvinski, L. & Li, Z. (2026). StableAML: Machine Learning for Behavioral Wallet Detection in Stablecoin Anti-Money Laundering on Ethereum. arXiv:2602.17842.
"""


def create_english_doc():
    """Create the English version."""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for i in range(5):
        hs = doc.styles[f'Heading {i+1}'] if i > 0 else doc.styles['Title']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    lines = EN_LINES.strip().split('\n')
    build_document(doc, lines)

    path = '/Users/tangliam/CriptoAnalyst/midterm_report_EN.docx'
    doc.save(path)
    print(f'English version saved: {path}')
    return path


def main():
    # Chinese version
    with open('/Users/tangliam/CriptoAnalyst/midterm_report_background.md', 'r', encoding='utf-8') as f:
        cn_lines = f.read().strip().split('\n')
    create_chinese_doc(cn_lines)

    # English version
    create_english_doc()

    print('\nDone! Both documents generated.')


if __name__ == '__main__':
    main()
