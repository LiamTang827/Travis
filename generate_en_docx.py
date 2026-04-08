"""
Generate IEEE-formatted English Word document for the 1st Interim Report.
- Two-column layout (simulated via IEEE-style formatting)
- Real experiment figures from ml/data/model_output/
- Proper Figure and Table numbering with captions
"""
import re, json
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IMG_DIR = '/Users/tangliam/CriptoAnalyst/ml/data/model_output'
OUT_PATH = '/Users/tangliam/CriptoAnalyst/midterm_report_EN.docx'


# ── helpers ──

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders (IEEE style: top/bottom only)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), val.get('sz', '4'))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            borders.append(el)
        else:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'none')
            el.set(qn('w:sz'), '0')
            el.set(qn('w:space'), '0')
            borders.append(el)
    tcPr.append(borders)


def parse_inline(paragraph, text, size=10, font='Times New Roman'):
    """Parse **bold** and `code` inline formatting."""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(size)
            run.font.name = font
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.size = Pt(size - 1)
            run.font.name = 'Courier New'
        elif part:
            run = paragraph.add_run(part)
            run.font.size = Pt(size)
            run.font.name = font


def add_caption(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER):
    """Add a centered caption paragraph (for Figure/Table labels)."""
    p = doc.add_paragraph()
    p.alignment = align
    # Split to bold the "Figure X." or "TABLE X" part
    m = re.match(r'^((?:Fig\.|Figure|TABLE|Table)\s+\w+[.:]?\s*)', text)
    if m:
        run = p.add_run(m.group(1))
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        rest = text[m.end():]
        if rest:
            run2 = p.add_run(rest)
            run2.font.size = Pt(9)
            run2.font.name = 'Times New Roman'
    else:
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(8)
    return p


def add_ieee_table(doc, headers, rows, caption_text):
    """Add an IEEE-style table with top/bottom rules and caption above."""
    # Caption above table
    add_caption(doc, caption_text)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove all default borders, set only IEEE horizontal rules
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        borders.append(el)
    # Remove existing borders if any
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

    # Header row
    border_thick = {'val': 'single', 'sz': '12', 'color': '000000'}
    border_thin = {'val': 'single', 'sz': '6', 'color': '000000'}

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        set_cell_borders(cell, top=border_thick, bottom=border_thin)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            parse_inline(p, val.strip(), size=9)
            if r_idx == len(rows) - 1:
                set_cell_borders(cell, bottom=border_thick)

    doc.add_paragraph('')
    return table


def add_figure(doc, img_path, caption_text, width=Inches(4.5)):
    """Add a figure with caption below."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=width)
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)

    add_caption(doc, caption_text)


def add_para(doc, text, size=10, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             first_indent=Cm(0.75), space_before=0, space_after=4):
    """Add a paragraph with IEEE styling."""
    p = doc.add_paragraph()
    p.alignment = align
    if first_indent:
        p.paragraph_format.first_line_indent = first_indent
    parse_inline(p, text, size=size)
    if bold:
        for run in p.runs:
            run.bold = True
    if italic:
        for run in p.runs:
            run.italic = True
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0
    return p


def add_heading_ieee(doc, text, level=1):
    """Add IEEE-style heading (centered for level 1, left for level 2+)."""
    p = doc.add_paragraph()
    if level == 0:
        # Title
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(24)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(4)
    elif level == 1:
        # Section heading: Roman numeral, centered, caps
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        # Subsection: italic, left-aligned
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
    return p


def add_bullet(doc, text, size=10):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.text = ''
    parse_inline(p, text, size=size)
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.0
    return p


def add_blockquote(doc, text, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.5)
    parse_inline(p, text, size=size)
    for run in p.runs:
        run.italic = True
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.0
    return p


# ── Table / Figure counters ──
_tbl_num = 0
_fig_num = 0

def next_table():
    global _tbl_num
    _tbl_num += 1
    return _tbl_num

def next_fig():
    global _fig_num
    _fig_num += 1
    return _fig_num


# ── Main document builder ──

def build_report():
    global _tbl_num, _fig_num
    _tbl_num = 0
    _fig_num = 0

    doc = Document()

    # Page setup — IEEE: US Letter, narrow margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.78)
    section.right_margin = Cm(1.78)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.0

    # ── TITLE ──
    add_heading_ieee(doc, 'Cross-Chain Tracing-Based Cryptocurrency\nAnti-Money Laundering Risk Identification System', level=0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('1st Interim Report')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(12)

    # ════════════════════════════════════════
    # I. INTRODUCTION
    # ════════════════════════════════════════
    add_heading_ieee(doc, 'I. Introduction', level=1)

    add_heading_ieee(doc, 'A. Cryptocurrency Money Laundering at Scale', level=2)

    add_para(doc, 'The decentralized, pseudonymous, and globally accessible nature of blockchain has made it a significant channel for illicit fund flows alongside financial innovation. According to Chainalysis\u2019s 2024 Crypto Money Laundering Report [10], illegal cryptocurrency addresses received approximately **$40.9 billion** in 2023, a figure that grew to over **$154 billion** by 2025 \u2014 a 162% year-over-year increase. Particularly noteworthy is that stablecoins have risen from a minor fraction to **63%** of illicit transactions (reaching 84% in 2025), as criminals increasingly shift from Bitcoin to stablecoins like USDT \u2014 driven by their superior liquidity, ease of cross-chain transfer, and larger regulatory blind spots.')

    add_para(doc, 'From a crime typology perspective, money laundering techniques have exhibited a systematic trend toward \u201cprofessionalization.\u201d The Huione Group, for example, has processed over **$70 billion** in cryptocurrency transactions since 2021, gradually evolving into an underground financial infrastructure serving the entire fraud and laundering pipeline. This trend indicates that cryptocurrency crime is no longer scattered individual behavior but organized crime with institutional structure and technical barriers to entry.')

    add_heading_ieee(doc, 'B. The Rise of Cross-Chain Bridges and Regulatory Blind Spots', level=2)

    add_para(doc, 'Cross-chain bridges have expanded dramatically in recent years. Stargate Finance alone handles over **$2.3 billion** in monthly cross-chain volume, with the broader DeFi ecosystem facilitating over **$8 billion** in monthly cross-chain asset transfers. While bridges serve a legitimate and critical infrastructure function \u2014 enabling asset transfers between different blockchains \u2014 this capability has been systematically exploited by criminals to sever fund tracing chains.')

    add_para(doc, 'Elliptic\u2019s 2023 report [11] identified **$7 billion** in illicit assets laundered through cross-chain services, with the figure growing rapidly since 2022. Among identifiable laundering schemes, **58% utilized cross-chain bridges** as a key component (2024 data). Chainalysis [10] similarly reported a significant surge in bridge usage from stolen-fund-linked addresses in 2023.')

    add_para(doc, 'The operations of North Korean hacking group Lazarus Group illustrate this pattern with striking clarity. In the **Ronin Bridge Attack (March 2022)**, approximately $625 million was stolen from the Ronin sidechain used by the play-to-earn game Axie Infinity. The stolen assets were subsequently laundered through Tornado Cash mixing, transferred via Avalanche cross-chain bridge to the Bitcoin network, and then processed through the Sinbad mixer for secondary washing. The entire laundering flow spanned over 12,000 addresses and involved multiple blockchains, making it one of the most complex laundering operations ever documented in the cryptocurrency space.')

    add_para(doc, 'In the **Harmony Horizon Bridge Attack (June 2022)**, approximately $100 million was stolen, with **98% of the stolen assets processed through Tornado Cash**. The funds were then repeatedly hopped between Ethereum, BNB Chain, and BitTorrent Chain, until partial funds resurfaced on Avalanche and TRON chains in 2023. The multi-chain, multi-hop strategy was specifically designed to exhaust the capacity of any single-chain tracing tool, forcing investigators to manually correlate events across disparate blockchain explorers.')

    add_para(doc, 'These attacks demonstrate the typical modern laundering pattern: **mixing + cross-chain + multi-hop relay**, the explicit purpose of which is to consume law enforcement and compliance resources through compounding tracing difficulty across chains and protocols. The fact that law enforcement recovered only approximately **$30 million** from the Ronin hack \u2014 a mere 4.8% of the stolen total \u2014 underscores the inadequacy of current tracing infrastructure.')

    add_heading_ieee(doc, 'C. Limitations of Existing Tools', level=2)

    add_para(doc, 'Investigations by INTERPOL and related law enforcement agencies reveal that **74% of agencies report significant limitations** in existing blockchain investigation tools when it comes to cross-chain activity tracing. Mainstream commercial tools (such as Chainalysis Reactor and TRM Labs) possess certain cross-chain capabilities, but their core algorithms have not been publicly validated by academia, and they primarily serve commercial institutional clients at price points inaccessible to ordinary users or small research teams.')

    add_para(doc, 'From the academic research perspective, existing literature suffers from four main limitations. First, **single-chain focus**: the vast majority of AML research targets Bitcoin or Ethereum as a single study object, lacking analytical frameworks for cross-chain scenarios. When funds move from Ethereum through a bridge to Arbitrum or Tron, existing academic tools simply lose the trail. Second, **insufficient blacklist coverage**: most academic studies rely solely on the OFAC sanctions list, overlooking the actual freeze lists maintained by stablecoin issuers such as Tether \u2014 whose USDT blacklist contains over 8,500 frozen addresses with assets exceeding $4.2 billion, providing a more granular and timely picture of on-chain enforcement than OFAC\u2019s broader but less blockchain-specific sanctions.')

    add_para(doc, 'Third, **no traceability classification for bridges**: existing research rarely distinguishes between \u201ctransparent bridges\u201d (where the counterpart address on the target chain can be determined through protocol APIs or event logs) and \u201copaque bridges\u201d (where the link between input and output is severed by liquidity pools or market-maker models). This classification is critical because it determines whether fund flow tracing can continue across chains or must terminate at the bridge as a dead end. Fourth, **lack of depth control in graph analysis**: current graph analysis methods typically set fixed depths (e.g., 5 hops) without considering the varying suspicion levels across different branches, leading to either too-shallow tracing (easily evadable by inserting additional relay addresses) or too-deep tracing (exponential computational cost and false positives on legitimate users who happen to be distant neighbors of illicit addresses).')

    add_heading_ieee(doc, 'D. The Passive Taint Problem for Ordinary Users', level=2)

    add_para(doc, 'Existing AML tools (Chainalysis, TRM Labs, Elliptic) are designed exclusively for institutional clients \u2014 exchanges and regulators. However, **ordinary on-chain users** face compliance risks that are often passive, unforeseeable, and entirely without recourse. These users have no access to the commercial risk-scoring infrastructure that exchanges rely upon, yet they bear the consequences when tainted funds flow through their addresses.')

    add_para(doc, 'Consider two representative scenarios. In the first scenario (incoming transfer taint), a blacklisted address A transfers funds to an ordinary user B, who subsequently deposits into a centralized exchange. The exchange\u2019s AML system detects that B\u2019s deposit originates from high-risk address A, and B\u2019s account is frozen or flagged as high-risk \u2014 despite B having no knowledge of A\u2019s blacklisted status. In the second, more insidious scenario (cross-chain taint propagation), blacklisted address A sends funds through a mixer, then a cross-chain bridge, then a relay address C, and finally to ordinary user B. User B has absolutely no way of knowing that the upstream of the transfer from C contains a blacklisted address. Yet the exchange\u2019s 2-hop or 3-hop backward scan may still trigger a risk alert on B\u2019s deposit.')

    add_para(doc, 'This phenomenon is known as **\u201cPassive Taint\u201d** or **\u201cInnocent Third-Party Harm.\u201d** The core contradiction is threefold. First, while blockchain is publicly transparent and anyone can theoretically verify an address\u2019s history, **analytical capabilities are monopolized by commercial companies**, leaving individual users unable to perform equivalent risk self-checks before initiating transactions. Second, exchange compliance policies are **opaque and inconsistent**: some exchanges trace 2 hops backward, others trace 5 hops, and users have no idea how deep their deposits will be scrutinized. Third, launderers intentionally **distribute funds to numerous ordinary addresses** (the Money Mule pattern), unknowingly enrolling these users as participants in the laundering chain, which then deposit into exchanges to bypass direct blacklist detection.')

    add_para(doc, 'BlockSec\u2019s analysis in Following the Frozen [12] found that **54% of blacklisted addresses had already transferred their assets before the freeze occurred** \u2014 meaning substantial \u201ctainted funds\u201d are already circulating among normal user addresses, while these users remain completely unaware of the contamination. This finding directly motivates the need for a self-service risk assessment tool that ordinary users can employ before transacting.')

    add_heading_ieee(doc, 'E. Research Objectives', level=2)

    add_para(doc, 'Synthesizing the above background, this research addresses the following core question:')

    add_blockquote(doc, 'Given a target blockchain address, how can we systematically identify whether it has fund associations with known blacklisted addresses, and what is the risk level and confidence of such associations \u2014 particularly when fund paths cross multiple blockchains, pass through mixers, or involve multi-layer relays?')

    add_para(doc, 'This core question decomposes into four specific sub-questions. Sub-question 1 (Classification): How to distinguish genuinely untraceable concealment behaviors (mixers, opaque bridges) from penetratable cross-chain behaviors (transparent bridges)? Sub-question 2 (Depth): In multi-hop tracing, how to find a reasonable balance between insufficient tracing depth (easily evadable) and excessive tracing overhead (false positives on innocent users)? Sub-question 3 (Scoring): How to quantify multi-hop, multi-chain tracing results into interpretable risk scores that enable horizontal comparison across different addresses? Sub-question 4 (Proportion): When an address holds both legitimate and tainted funds simultaneously, how to distinguish the proportion of \u201cdirty money\u201d from \u201cclean money,\u201d avoiding blanket high-risk classification of the entire address?')

    add_heading_ieee(doc, 'F. Potential Outcome', level=2)

    add_para(doc, 'The expected deliverable is a **prototype software system** consisting of four components: (1) a single-address risk analysis engine supporting Ethereum and Tron chains, with integrated blacklist detection (8,500+ USDT frozen addresses), cross-chain bridge registry (20+ contracts distinguishing transparent and opaque bridges), and mixer identification (10 Tornado Cash contracts); (2) a recursive BFS trace graph engine with cross-chain bridge resolution via LayerZero Scan API, adaptive depth control, and convergence tracking for scatter-then-converge patterns; (3) a machine learning module for behavioral wallet classification using 61 domain-specific features and tree ensemble models; and (4) visualization outputs in both JSON format (for programmatic consumption) and Mermaid graph format (for human review and presentation).')

    add_para(doc, 'The system targets **ordinary on-chain users** \u2014 not institutions \u2014 filling a critical tool gap identified by the literature. No open-source, self-service address risk assessment tool currently exists that combines multi-chain tracing with ML-augmented scoring. Users will be able to check whether a counterparty address has fund associations with known violations before initiating transactions or accepting transfers, thereby proactively avoiding passive taint risk.')

    # ════════════════════════════════════════
    # II. RELATED WORK
    # ════════════════════════════════════════
    add_heading_ieee(doc, 'II. Related Work', level=1)

    add_heading_ieee(doc, 'A. Blockchain Transaction Graphs and AML Detection', level=2)

    add_para(doc, 'Modeling blockchain transactions as graphs is the mainstream paradigm in current AML research. The most influential foundational work in this direction comes from Weber et al. (2019), who released the **Elliptic Dataset** \u2014 a Bitcoin transaction graph containing 203,769 nodes and 234,355 directed edges, with approximately 4,500 nodes labeled as \u201cillicit\u201d \u2014 and first proposed applying Graph Convolutional Networks (GCN) to AML classification at the KDD 2019 Anomaly Detection in Finance Workshop [1]. This dataset remains the most widely used AML benchmark in academia to this day. In 2024, the same team released the second-generation dataset **Elliptic2**, providing community-level annotations to support morphological analysis of entire laundering subgraphs rather than individual transactions [2].')

    add_para(doc, 'However, these studies share a critical limitation: **all are based on single chains (primarily Bitcoin or Ethereum), without discussing how to continue tracing after funds cross chains.** Additionally, address clustering \u2014 a core capability of commercial tools (Chainalysis, Elliptic) that consolidates multiple addresses controlled by the same entity \u2014 remains algorithmically opaque; academic open-source implementations rely primarily on heuristic methods with limited coverage under Ethereum\u2019s account model. This research does not directly implement address clustering but uses interaction frequency as a proxy metric during child node generation, prioritizing high-frequency counterparties to partially compensate for this gap.')

    add_heading_ieee(doc, 'B. Cross-Chain Transaction Tracing and Traceability', level=2)

    add_para(doc, 'This direction provides the most direct academic context for this research and is one of the fastest-growing research areas in recent years. **Mazorra et al. (2023/2024)** in Tracing Cross-chain Transactions between EVM-based Blockchains (published in the Ledger journal) proposed a set of heuristic matching algorithms for cross-chain transactions between EVM-compatible chains [3]. The core insight is: between EVM-compatible chains, user addresses remain consistent across different chains (the same `0xABCD...` address on Ethereum and Polygon is controlled by the same private key), enabling matching of Lock events on the source chain with Mint/Release events on the target chain through a \u201ctime window + amount + token type\u201d combination algorithm. This research achieved a **99.65%** deposit matching rate and **92.78%** withdrawal matching rate on over 2 million cross-chain transactions spanning August 2020 to August 2023.')

    add_para(doc, '**Sun et al. (2025)** in Track and Trace: Automatically Uncovering Cross-chain Transactions in the Multi-blockchain Ecosystems (arXiv 2504.01822) systematically analyzed **12 major cross-chain bridges** (including Stargate, Celer cBridge, Wormhole, Synapse, etc.), covering Ethereum source chain data from April 2021 to March 2024, and proposed a general framework for automatically identifying cross-chain transactions [4]. A key finding was that transparency varies dramatically across bridges \u2014 message-passing protocol-based bridges (e.g., LayerZero) can obtain counterpart transaction hashes directly via API, while liquidity pool-based bridges (e.g., Synapse) make it nearly impossible to find explicit input-output correspondences in on-chain data.')

    add_para(doc, 'A Survey of Transaction Tracing Techniques for Blockchain Systems (arXiv 2510.09624) provides a macro-level overview of blockchain transaction tracing techniques, categorizing existing methods into four classes: on-chain event correlation, API-assisted tracing, statistical inference, and machine learning. The survey identifies cross-chain tracing as the area most lacking systematic research [5].')

    add_para(doc, 'A critical distinction \u2014 not yet adequately discussed in the literature \u2014 is between **transparent bridges** (where the counterpart address can be determined) and **opaque bridges** (where the input-output link is severed). Table I formalizes this taxonomy, which this research treats as a first-class design decision.')

    # TABLE I: Bridge classification
    tn = next_table()
    add_ieee_table(doc,
        ['Dimension', 'Transparent Bridge', 'Opaque Bridge'],
        [
            ['Mechanism', 'Message-passing (LayerZero, Wormhole) or Rollup official bridges', 'Liquidity pools (Synapse), market-maker (Orbiter, Owlto)'],
            ['On-chain linkage', 'Unique transferId or shared txHash links both ends', 'Funds enter shared pool; no linkage possible'],
            ['Laundering risk', 'Low (traceable)', 'High (equivalent to mixer)'],
            ['Examples', 'Stargate, Hop, Arbitrum/Optimism bridges', 'Multichain (collapsed), Orbiter, Synapse'],
        ],
        f'TABLE {tn}: Transparent vs. opaque cross-chain bridge classification (synthesized from [3], [4], [5]).'
    )

    add_para(doc, 'Multichain (formerly one of the largest cross-chain bridges) collapsed in 2023 due to internal issues, resulting in approximately **$127 million** in lost assets \u2014 an event that itself exposed the fundamental deficiencies of opaque bridges in transparency and auditability. The laundering paths used by Lazarus Group in the Ronin and Harmony attacks (Section I-B) further illustrate that the combination of mixers and multi-chain hops pushes recovery rates to extremely low levels (4.8% for Ronin), underscoring the need for better automated tracing tools.')

    add_heading_ieee(doc, 'C. Multi-Hop Risk Scoring and Tree-Based Tracing', level=2)

    add_para(doc, '**Moser, Bohme & Breuker (2014)** in Towards Risk Scoring of Bitcoin Transactions (Financial Cryptography 2014) first formalized \u201cpropagating risk from known illicit addresses along the transaction graph\u201d as a research problem [6]. Two propagation strategies were proposed: **Poison** (full taint) \u2014 if any illicit input exists in the fund sources, the output is considered fully tainted; and **Haircut** (proportional) \u2014 taint proportion equals the ratio of illicit source funds to total inputs, gradually diluted through mixing. This is the earliest academic work discussing \u201cmulti-hop risk\u201d and has been cited by virtually all subsequent Taint Analysis research. Its limitation lies in its 2014 publication date, with the scope limited to Bitcoin single-chain.')

    add_para(doc, '**Hercog & Povsea (2019)** proposed the **TaintRank** algorithm [7], analogizing taint propagation to PageRank: constructing a weighted directed graph with addresses as nodes and transactions as edges, where each node\u2019s taint value is the weighted sum of all upstream nodes\u2019 taint values. Taint naturally decays with propagation distance, with the final distribution following a power-law shape. TaintRank scores the entire Bitcoin network in a batch, global manner, producing a 0\u20131 taint index for each address. Unlike this research, it is an offline batch-processing algorithm that does not support real-time tree-based queries.')

    add_para(doc, '**Liao, Zeng, Belenkiy & Hirshman (2025)** from USDC issuer Circle, in Transaction Proximity: A Graph-Based Approach to Blockchain Fraud Prevention (arXiv:2505.24284), applied the BFS approach to the entire Ethereum historical graph [8]: **206 million nodes, 442 million edges**, covering all transactions from genesis to May 2024. With a 5-hop BFS limit (covering 98.2% of active USDC holders), they introduced **Transaction Proximity** (shortest hop count to a regulated exchange) and **EAI** (Easily Attainable Identities \u2014 addresses directly connected to exchanges). A key finding: 83% of known attacker addresses are not EAI, and 21% are more than 5 hops from any regulated exchange, demonstrating that criminal addresses tend to be structurally distant from normal circulation nodes.')

    add_para(doc, 'Notably, Circle\u2019s risk perspective is **opposite but complementary** to ours: Transaction Proximity measures distance from legitimate anchors (closer = more legitimate), while this research measures distance from illicit anchors (closer = more dangerous). The two methods can theoretically be fused to provide confidence from both directions for the same address. Table II provides a systematic comparison across all discussed approaches.')

    # TABLE II: Systematic comparison
    tn = next_table()
    add_ieee_table(doc,
        ['Feature', 'Moser 2014', 'TaintRank 2019', 'Tx Proximity 2025', 'Chainalysis', 'This Research'],
        [
            ['Cross-chain tracing', 'No', 'No', 'No', 'Partial', 'Yes'],
            ['Bridge classification', 'No', 'No', 'No', 'No', 'Yes'],
            ['Adaptive depth', 'No', 'No', 'No (fixed 5)', 'No', 'Yes'],
            ['Real-time query', 'No', 'No (batch)', 'No (offline)', 'Yes', 'Yes'],
            ['For ordinary users', '\u2014', '\u2014', '\u2014', 'No', 'Yes'],
            ['USDT freeze list anchor', 'No', 'No', 'No', 'Partial', 'Yes'],
            ['Explicit hop decay', 'Yes', 'Implicit', 'No', 'N/A', 'Yes (x0.6)'],
            ['Open-source', '\u2014', 'Partial', 'Partial', 'No', 'Yes'],
        ],
        f'TABLE {tn}: Systematic comparison with existing approaches (compiled from [6], [7], [8]).'
    )

    add_heading_ieee(doc, 'D. Machine Learning for Stablecoin AML', level=2)

    add_para(doc, '**Juvinski & Li (2026)** conducted the most comprehensive ML-based study on stablecoin AML to date in the **StableAML** paper [24]. They trained tree ensemble models on 16,433 labeled addresses using 68 behavioral features organized into four categories: interaction features (contacts with known entities), transfer features (amount/volume statistics), network features (graph topology metrics), and temporal features (timing patterns). CatBoost achieved the best performance with Macro-F1 = 0.9775, significantly outperforming graph neural networks (GraphSAGE, F1 = 0.8048).')

    add_para(doc, 'Their most significant finding was: **domain-specific feature engineering matters more than complex graph algorithms** on stablecoin transaction graphs. The reason is fundamental: stablecoin graphs are extremely sparse (density < 0.01), meaning most addresses have very few direct connections. In such sparse graphs, GNN message passing cannot effectively propagate information because there are too few neighbors to aggregate over. Tree ensemble models, by contrast, can directly leverage hand-crafted features that encode domain knowledge about laundering patterns without relying on graph connectivity. This finding has direct implications for this research: rather than investing in computationally expensive GNN architectures, a feature engineering + tree ensemble approach is more likely to yield practical results, especially given our smaller labeled dataset.')

    add_heading_ieee(doc, 'E. Regulatory Frameworks and Real-World Needs', level=2)

    add_para(doc, '**FATF (Financial Action Task Force)** incorporated virtual assets (VA) and virtual asset service providers (VASPs) into its AML/CFT standards framework (Recommendation 15) in 2019. In its 2023 targeted update report [9], FATF noted that among 151 member jurisdictions, **more than half had not yet implemented the \u201cTravel Rule,\u201d** and 75% were partially or non-compliant with R.15. The report also specifically highlighted the significant growth in stablecoin use by DPRK actors, terrorism financiers, and drug traffickers.')

    add_para(doc, 'The **EU AI Act (Regulation 2024/1689)**, which entered into force in 2024, classifies AI systems used for AML/CFT compliance as **\u201cHigh-Risk AI Systems\u201d (Annex III)** [18]. The Act requires such systems to satisfy three conditions: (1) sufficient transparency of decision logic, (2) human oversight mechanisms, and (3) the ability to provide complete decision explanations to regulators and end users. Violators face fines up to **EUR 35 million or 7% of global turnover**. This creates a core tension in the field: the models with the best detection performance (GNN, deep learning) are precisely the least interpretable, while the most interpretable methods (rule engines) have limited detection capability. This tension directly motivates the hybrid approach adopted in this research.')

    add_heading_ieee(doc, 'F. Interpretability and Privacy-Preserving Compliance', level=2)

    add_para(doc, 'The interpretability of AML systems is not optional but a regulatory mandate. FATF Recommendation 20 requires that Suspicious Transaction Reports (STRs) include textual explanations of why a transaction is considered suspicious. The US Bank Secrecy Act similarly requires natural language descriptions of suspicious behavioral patterns in SARs. The EU AI Act extends this obligation to AI systems themselves [18].')

    add_para(doc, 'Watson, Richards & Schiff (2025) proposed a representative three-layer architecture for explainable blockchain AML (architecture from [14]): Layer 1 uses a GNN detection layer (GCN-GRU hybrid, accuracy 0.9470, AUC-ROC 0.9807); Layer 2 uses GraphLIME attribution (identifying which features drove the classification); Layer 3 uses an LLM explanation layer (converting attributions into natural language). Nicholls et al. (2024) demonstrated that LLMs can independently generate actionable explanations from Bitcoin transaction data, and that these explanations themselves can serve as features for downstream detection [15]. Sun et al. (2024) published the first systematic review of LLM applications in blockchain security [16].')

    add_para(doc, 'On the privacy front, Buterin et al. (2023) introduced **Privacy Pools** [19], demonstrating that privacy and compliance need not be mutually exclusive \u2014 users can prove fund source legitimacy via zero-knowledge proofs without revealing transaction details. Privacy Pools v1 launched on Ethereum mainnet in 2024. However, Constantinides & Cartlidge (2025) identified that Proof of Innocence approaches depend on blacklist completeness and timeliness [20] \u2014 if a deposit is flagged as illicit after passing the PoI check, it has already entered the privacy pool irrevocably. The Tornado Cash sanctions saga (OFAC sanction in 2022, Fifth Circuit reversal in 2024, Treasury delisting in 2025 [21]) further illustrates that sanctioning privacy tools is not sustainable, and technical solutions for the privacy-compliance balance are needed.')

    add_para(doc, 'This research\u2019s rule engine (BFS tracing + risk scoring) has inherent interpretability advantages: every decision step (why a node was flagged, what the risk path is, how decay was calculated) has a complete causal chain with no \u201cblack box\u201d problem. This makes introducing LLMs as an explanation layer straightforward as a future extension \u2014 directly converting structured JSON trace results into natural language risk reports without requiring post-hoc attribution tools like SHAP or LIME.')

    add_heading_ieee(doc, 'G. Summary: Advantages and Disadvantages of Existing Solutions', level=2)

    add_para(doc, '**Commercial tools** (Chainalysis, TRM Labs, Elliptic) offer broad coverage and some cross-chain capability, but are closed-source, expensive, institution-only, and algorithmically unvalidated by academia. Their \u201cblack box\u201d nature increasingly conflicts with EU AI Act requirements for decision transparency. **Academic graph-based approaches** (Elliptic/GNN, TaintRank, Transaction Proximity) are reproducible and theoretically grounded, but limited to single chains, batch/offline processing, and institutional research use cases. None supports real-time single-address queries targeting ordinary users.')

    add_para(doc, '**This research** bridges the gap by combining four elements absent in prior work: (1) cross-chain tracing with explicit bridge traceability classification \u2014 absent in all prior academic work; (2) adaptive-depth BFS \u2014 a data-driven alternative to fixed-depth scanning that traces deeper only where suspicion warrants it; (3) ML-augmented hybrid scoring \u2014 complementing the rule engine with behavioral pattern recognition following the StableAML methodology; and (4) open-source, individual-user focus \u2014 filling the tool gap for ordinary on-chain users who currently have no self-service risk assessment capability.')

    # ════════════════════════════════════════
    # III. SYSTEM MODELING AND STRUCTURE
    # ════════════════════════════════════════
    add_heading_ieee(doc, 'III. System Modeling and Structure', level=1)

    add_heading_ieee(doc, 'A. Architecture Overview', level=2)

    add_para(doc, 'The system consists of three core modules: (1) **Single-Address Analysis Engine** (`aml_analyzer.py`) \u2014 fetches on-chain data via Etherscan/TronScan APIs, checks against blacklists, bridge registry, and mixer contracts, performs ML-augmented risk scoring; (2) **Recursive Trace Graph Engine** (`trace_graph.py`) \u2014 BFS-based multi-hop fund association tree with cross-chain bridge resolution, adaptive depth, and convergence tracking; (3) **ML Risk Scorer** (`ml/` pipeline) \u2014 behavioral feature extraction, tree ensemble classification, integrated as hybrid scoring component.')

    add_heading_ieee(doc, 'B. Design Choices and Justifications', level=2)

    add_para(doc, '**Bridge traceability classification as first-class citizen.** Unlike existing research, this system classifies each bridge as transparent or opaque. Transparent bridges are resolved via protocol APIs for continued tracing; opaque bridges are treated as mixer-equivalent terminal nodes. This addresses the traceability disparity identified by Sun et al. [4].')

    add_para(doc, '**Adaptive depth BFS.** Normal branches use standard maximum depth (default 3 hops); suspicious branches (blacklist contact, mixer, opaque bridge) receive extra hops (`depth_bonus`, default +1). This provides a data-driven balance between thoroughness and cost.')

    add_para(doc, '**Bidirectional entity detection.** Code review revealed that checking only the `to` field misses funds extracted from mixers (`from=mixer`). The system checks both `from` and `to` against all entity registries, recording direction (IN/OUT).')

    add_para(doc, '**Convergence tracking.** The original BFS silently discarded convergence paths. The redesigned system preserves `converge_from` and `in_degree` metadata without re-expanding, capturing scatter-then-converge laundering patterns.')

    add_para(doc, '**Hybrid rule + ML scoring.** Pure rule engines have hard-coded blind spots that sophisticated launderers can learn to exploit; pure ML models may miss known high-severity signals such as direct mixer contact or blacklist adjacency. The hybrid formula `risk_score = rule_engine x 0.4 + ML_predict_proba x 0.6` addresses both weaknesses. The rule engine retains 40% weight to ensure that known high-risk signals (mixer contact, direct blacklist association) are never underestimated regardless of the ML model\u2019s output. The ML model contributes 60% for behavioral patterns \u2014 temporal anomalies, amount distributions, network topology signatures \u2014 that hard-coded rules cannot capture. The system supports graceful degradation: if the trained model files are absent or corrupted, it automatically falls back to pure rule-engine scoring without any user intervention.')

    # TABLE III: Addressing limitations
    tn = next_table()
    add_ieee_table(doc,
        ['Existing Limitation', 'How This System Addresses It'],
        [
            ['Single-chain only', 'Multi-chain tracing via bridge registry + LayerZero Scan API'],
            ['No bridge classification', 'Explicit transparent/opaque taxonomy driving tracing logic'],
            ['Fixed depth', 'Adaptive depth with suspicion-triggered bonus'],
            ['Institution-only tools', 'Open-source, designed for individual user self-assessment'],
            ['Black-box models', 'Rule engine with full causal chain; ML with feature importance'],
            ['OFAC-only blacklists', 'USDT freeze list (8,500+ addresses) as primary anchor'],
        ],
        f'TABLE {tn}: How this system addresses existing limitations.'
    )

    # ════════════════════════════════════════
    # IV. METHODOLOGY AND ALGORITHMS
    # ════════════════════════════════════════
    add_heading_ieee(doc, 'IV. Methodology and Algorithms', level=1)

    add_heading_ieee(doc, 'A. Risk Scoring Algorithm', level=2)

    add_para(doc, 'The hop-decay risk scoring applies a **0.6x per-hop** decay coefficient, reflecting the consensus that direct association carries higher risk than indirect association. Table IV shows the effective risk at each hop distance.')

    tn = next_table()
    add_ieee_table(doc,
        ['Association Distance', 'Effective Risk (base 100)', 'Risk Level'],
        [
            ['Direct contact (1 hop)', '60', 'HIGH'],
            ['2 hops', '36', 'MEDIUM'],
            ['3 hops', '21.6', 'LOW'],
            ['4+ hops', '\u2264 13', 'Reference only'],
        ],
        f'TABLE {tn}: Hop-decay risk scoring.'
    )

    add_para(doc, 'After BFS expansion, a secondary scan identifies **suspect relay addresses** (Money Mules): ostensibly clean addresses whose subtrees contain blacklist hits receive a subtree contamination score.')

    add_heading_ieee(doc, 'B. Sampling Bias Corrections', level=2)

    add_para(doc, 'Five systematic bugs were identified through code review. **Sampling bias** (P1, P2, P5): P1 \u2014 `MAX_TX_FETCH=100` truncated histories, fixed by increasing to 500; P2 \u2014 `txlist+tokentx` double-counted transactions, fixed by `(hash, from, to)` deduplication; P5 \u2014 frequency-only counterparty ranking let DEX routers dominate, fixed by amount-weighted composite scoring. **Graph distortion** (P3, P4): P3 \u2014 bridge/mixer detection only checked `to`, fixed by bidirectional detection; P4 \u2014 convergence paths silently discarded, fixed by preserving metadata. The core insight: **these bugs\' combined effect was that launderers\' adversarial strategies precisely exploited the system\'s blind spots**.')

    add_heading_ieee(doc, 'C. Machine Learning Pipeline', level=2)

    add_para(doc, 'Following StableAML\'s finding [24], the ML pipeline has four stages:')

    add_para(doc, '**Data Collection.** Blocklisted (100 addresses from USDT blacklist), Normal (50 randomly sampled active addresses), and Sanctioned (10 from OFAC SDN). For each address, Etherscan getLogs API fetches all USDT/USDC Transfer events bidirectionally (topic[1] for sent, topic[2] for received).')

    add_para(doc, '**Feature Engineering.** Following StableAML\'s four-category framework (adapted from [24]), 61 behavioral features were extracted, as shown in Table V.')

    tn = next_table()
    add_ieee_table(doc,
        ['Category', 'Count', 'Representative Features'],
        [
            ['Interaction', '18', 'sent_to_mixer, received_from_mixer, has_flagged_interaction'],
            ['Transfer', '19', 'transfers_over_10k, drain_ratio, repeated_amount_ratio'],
            ['Network', '10', 'in_degree, out_degree, counterparty_flagged_ratio'],
            ['Temporal', '8', 'has_daily_burst, rapid_tx_ratio, hour_concentration'],
        ],
        f'TABLE {tn}: 61 behavioral features in four categories (framework adapted from [24]).'
    )

    add_para(doc, '**Model Training.** 5-Fold Stratified Cross Validation was used to compare four tree ensemble models: RandomForest, XGBoost, LightGBM, and CatBoost. Class imbalance (100 blocklisted vs. 50 normal) was handled via balanced class weights (RandomForest, LightGBM), `scale_pos_weight` (XGBoost), or `auto_class_weights` (CatBoost). Stratified splitting ensured that each fold maintained the original class distribution, critical for reliable evaluation on a small dataset.')

    add_para(doc, '**Model Integration.** The best model is serialized as a `.pkl` file alongside a `model_meta.json` containing feature names, label encoding, and top feature importances. At runtime, the `MLRiskScorer` class loads these artifacts, extracts features from the existing `RiskReport` object (which already contains counterparty statistics, entity interaction flags, and amount summaries from the rule engine), calls `predict_proba` to obtain the probability of the \u201cblocklisted\u201d class, and scales it to a 0\u2013100 risk score. This score is then combined with the rule engine score in the hybrid formula described in Section III-B.')

    # ════════════════════════════════════════
    # V. PRELIMINARY PERFORMANCE ANALYSIS
    # ════════════════════════════════════════
    add_heading_ieee(doc, 'V. Preliminary Performance Analysis', level=1)

    add_heading_ieee(doc, 'A. Model Comparison', level=2)

    add_para(doc, 'Fig. 1 shows the Macro-F1 and PR-AUC scores across all four models. RandomForest achieved the best performance with Macro-F1 = 0.919 and PR-AUC = 0.949. Table VI provides detailed per-model results.')

    # Figure 1: Model comparison
    fn = next_fig()
    add_figure(doc, f'{IMG_DIR}/model_comparison.png',
               f'Fig. {fn}. Model comparison: Macro-F1 and PR-AUC across four tree ensemble models.',
               width=Inches(5.0))

    # TABLE VI: Model results
    tn = next_table()
    add_ieee_table(doc,
        ['Model', 'Macro-F1', 'PR-AUC', 'Accuracy'],
        [
            ['**RandomForest**', '**0.919**', '**0.949**', '**0.927**'],
            ['XGBoost', '0.886', '0.917', '0.900'],
            ['CatBoost', '0.872', '0.937', '0.887'],
            ['LightGBM', '0.865', '0.932', '0.880'],
        ],
        f'TABLE {tn}: 5-Fold Stratified Cross Validation results for four tree ensemble models.'
    )

    add_heading_ieee(doc, 'B. Confusion Matrix Analysis', level=2)

    add_para(doc, 'Fig. 2 shows the confusion matrix for the best model (RandomForest). The model achieves 95.9% precision and 93.0% recall for blocklisted detection, with only 4 false positives (normal addresses incorrectly flagged as blocklisted) and 7 false negatives. Given the small training set (150 samples vs. StableAML\'s 16,433), these results are encouraging.')

    # Figure 2: Confusion matrix
    fn = next_fig()
    add_figure(doc, f'{IMG_DIR}/confusion_matrix.png',
               f'Fig. {fn}. Confusion matrix for the best model (RandomForest). TN=93, FP=7, FN=4, TP=46.',
               width=Inches(3.5))

    add_heading_ieee(doc, 'C. Feature Importance', level=2)

    add_para(doc, 'Fig. 3 shows the top 20 features by Mean Decrease Impurity (MDI) for RandomForest. The most discriminative feature is `drain_ratio` (balance drain rate) \u2014 blocklisted addresses have a mean of 0.21 vs. 1.37 for normal addresses, reflecting that frozen addresses cannot transfer funds out. Other top features include `out_degree`, `sent_count`, `unique_receivers`, and `in_out_degree_ratio`.')

    # Figure 3: Feature importance
    fn = next_fig()
    add_figure(doc, f'{IMG_DIR}/importance_randomforest.png',
               f'Fig. {fn}. Top 20 feature importances (MDI) for the RandomForest model.',
               width=Inches(5.0))

    add_para(doc, 'Fig. 4 shows the feature importance for XGBoost (by Gain). The ranking differs: `total_sent_amount` and `counterparty_flagged_ratio` are the top two, reflecting that XGBoost leverages KYC label data and large transaction amounts more heavily than RandomForest.')

    # Figure 4: XGBoost importance
    fn = next_fig()
    add_figure(doc, f'{IMG_DIR}/importance_xgboost.png',
               f'Fig. {fn}. Top 20 feature importances (Gain) for the XGBoost model.',
               width=Inches(5.0))

    add_heading_ieee(doc, 'D. System Module Status', level=2)

    tn = next_table()
    add_ieee_table(doc,
        ['Module', 'Status', 'Description'],
        [
            ['aml_analyzer.py', 'Implemented', 'Etherscan + TronScan API, blacklist (8,500+), bridge registry (20+), mixer (10), hybrid scoring'],
            ['trace_graph.py', 'Implemented', 'BFS tree, cross-chain resolution, adaptive depth, convergence tracking, Mermaid export'],
            ['ml/ pipeline', 'Implemented', 'Data collection, 61 features, 4 models trained, best model integrated'],
            ['Bug fixes (P1-P5)', 'Implemented', 'Sampling bias and graph distortion corrections'],
        ],
        f'TABLE {tn}: Current implementation status of system modules.'
    )

    add_heading_ieee(doc, 'E. Known Limitations', level=2)

    add_para(doc, 'Five limitations are identified. First, **limited data**: 150 samples is far fewer than StableAML\u2019s 16,433, and expanding the dataset is the most direct path to improved generalization. Second, **unverified Normal class**: randomly sampled \u201cnormal\u201d addresses from the chain may include unlabeled laundering addresses (label noise), which could depress precision. Third, **USDT/USDC only**: launderers who swap to ETH or other tokens via DEX protocols escape the current analysis scope entirely. Fourth, **real-time feature gap**: the ML scorer currently extracts features from the `RiskReport` object, which is a subset of the full behavioral profile; some temporal features (e.g., burst detection, interval statistics) require complete Transfer event data not available in the real-time scoring path. Fifth, **binary classification**: the current model only distinguishes blocklisted from normal, without finer sub-categories such as sanctioned, cybercrime, or terrorism financing \u2014 categories that would require additional labeled data sources to train.')

    # ════════════════════════════════════════
    # VI. MILESTONES AND SCHEDULE
    # ════════════════════════════════════════
    add_heading_ieee(doc, 'VI. Milestones and Overall Schedule', level=1)

    tn = next_table()
    add_ieee_table(doc,
        ['Phase', 'Period', 'Milestone'],
        [
            ['1', 'Weeks 1\u20133 (Done)', 'Core system: aml_analyzer.py, trace_graph.py, registries'],
            ['2', 'Weeks 4\u20135 (Done)', 'ML pipeline: data collection, features, training, integration'],
            ['3', 'Weeks 5\u20136 (Done)', 'Bug fixes P1\u2013P5, hybrid scoring, convergence tracking'],
            ['4', 'Weeks 7\u20139', 'End-to-end validation; expand dataset to 1,000+ addresses'],
            ['5', 'Weeks 10\u201312', 'Taint proportion analysis (FIFO/Haircut/Poison)'],
            ['6', 'Weeks 13\u201315', 'Additional bridge support; performance benchmarks'],
            ['7', 'Weeks 16\u201318', 'LLM explanation layer (exploratory); final report'],
        ],
        f'TABLE {tn}: Project milestones and schedule.'
    )

    # ════════════════════════════════════════
    # VII. WORK TO BE COMPLETED
    # ════════════════════════════════════════
    add_heading_ieee(doc, 'VII. Work to Be Completed for the Next Report', level=1)

    items = [
        '**Dataset expansion**: Scale from 150 to 1,000+ labeled addresses with manual verification of the Normal class.',
        '**End-to-end validation**: Run the complete system on known laundering cases (Ronin Bridge, Harmony Bridge) to verify trace tree reconstruction.',
        '**OFAC SDN integration**: Supplement the blacklist with sanctioned addresses not covered by Tether\'s freeze list.',
        '**Additional bridge support**: Implement counterpart address resolution for Hop Protocol and Across Protocol.',
        '**Time window filtering**: Introduce configurable time windows to exclude ancient transactions that introduce false positives.',
        '**Taint proportion analysis**: Implement FIFO and Haircut methods to upgrade from binary risk to proportional confidence scoring.',
        '**LLM explanation layer** (exploratory): Prototype feeding JSON trace results into an LLM to generate natural language risk narratives, as validated by [14] and [15].',
    ]
    for item in items:
        add_bullet(doc, item)

    # ════════════════════════════════════════
    # REFERENCES
    # ════════════════════════════════════════
    add_heading_ieee(doc, 'References', level=1)

    refs = [
        '[1] M. Weber et al., "Anti-Money Laundering in Bitcoin: Experimenting with Graph Convolutional Networks for Financial Forensics," KDD Workshop on Anomaly Detection in Finance, 2019. arXiv:1908.02591.',
        '[2] C. Bellei et al., "The Shape of Money Laundering: Subgraph Representation Learning on the Blockchain with the Elliptic2 Dataset," 2024. arXiv:2404.19109.',
        '[3] B. Mazorra et al., "Tracing Cross-Chain Transactions Between EVM-Based Blockchains: An Analysis of Ethereum-Polygon Bridges," Ledger Journal, 2023. arXiv:2504.15449.',
        '[4] X. Sun et al., "Track and Trace: Automatically Uncovering Cross-chain Transactions in the Multi-blockchain Ecosystems," 2025. arXiv:2504.01822.',
        '[5] J. Ren et al., "A Survey of Transaction Tracing Techniques for Blockchain Systems," 2025. arXiv:2510.09624.',
        '[6] M. Moser, R. Bohme, and D. Breuker, "Towards Risk Scoring of Bitcoin Transactions," Financial Cryptography Workshops, Springer, 2014.',
        '[7] U. Hercog and A. Povsea, "Taint Analysis of the Bitcoin Network," 2019. arXiv:1907.01538.',
        '[8] G. Liao, Z. Zeng, M. Belenkiy, and J. Hirshman, "Transaction Proximity: A Graph-Based Approach to Blockchain Fraud Prevention," Circle Research, 2025. arXiv:2505.24284.',
        '[9] FATF, "Targeted Update on Implementation of the FATF Standards on Virtual Assets and VASPs," June 2023.',
        '[10] Chainalysis, "2024 Crypto Money Laundering Report," Chainalysis Inc., 2024.',
        '[11] Elliptic, "$7 Billion in Crypto Laundered Through Cross-Chain Services," 2023.',
        '[12] BlockSec, "Following the Frozen: An On-Chain Analysis of USDT Blacklisting," 2023.',
        '[13] M. Weber and C. Bellei, "Elliptic Data Set," Kaggle / Elliptic, 2019.',
        '[14] A. Watson, G. Richards, and D. Schiff, "Explain First, Trust Later: LLM-Augmented Explanations for Graph-Based Crypto Anomaly Detection," 2025. arXiv:2506.14933.',
        '[15] J. Nicholls et al., "Large Language Model XAI Approach for Illicit Activity Investigation in Bitcoin," Neural Computing and Applications, Springer, 2024.',
        '[16] H. Sun et al., "Large Language Models for Blockchain Security: A Systematic Literature Review," 2024. arXiv:2403.14280.',
        '[17] D. Kute et al., "Explainable and Fair Anti-Money Laundering Models Using a Reproducible SHAP Framework," Discover Artificial Intelligence, Springer, 2026.',
        '[18] European Union, "Regulation (EU) 2024/1689 \u2014 AI Act, Annex III: High-Risk AI Systems," 2024.',
        '[19] V. Buterin et al., "Blockchain Privacy and Regulatory Compliance: Towards a Practical Equilibrium," Blockchain: Research and Applications, vol. 5, no. 1, 2023.',
        '[20] T. Constantinides and J. Cartlidge, "zkMixer: A Configurable Zero-Knowledge Mixer with AML Consensus Protocols," IEEE DAPPS, 2025. arXiv:2503.14729.',
        '[21] A. Brownworth et al., "Regulating Decentralized Systems: Evidence from Sanctions on Tornado Cash," NY Fed Staff Reports, No. 1112, 2024.',
        '[22] A. Chaudhary, "zkFi: Privacy-Preserving and Regulation Compliant Transactions using Zero Knowledge Proofs," 2023. arXiv:2307.00521.',
        '[23] F. Effendi and A. Chattopadhyay, "Privacy-Preserving Graph-Based ML with FHE for Collaborative AML," SPACE 2024. arXiv:2411.02926.',
        '[24] L. Juvinski and Z. Li, "StableAML: Machine Learning for Behavioral Wallet Detection in Stablecoin AML on Ethereum," 2026. arXiv:2602.17842.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(ref)
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.0
        pf.left_indent = Cm(0.5)
        pf.first_line_indent = Cm(-0.5)

    doc.save(OUT_PATH)
    print(f'IEEE-formatted English report saved: {OUT_PATH}')
    return OUT_PATH


if __name__ == '__main__':
    build_report()
